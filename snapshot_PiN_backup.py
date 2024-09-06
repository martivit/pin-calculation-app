import pandas as pd
from fuzzywuzzy import process
import numpy as np
import datetime
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Border, Side, Font, Alignment
from openpyxl.cell.cell import MergedCell
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
import matplotlib.pyplot as plt
from io import BytesIO
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib.text import Text

# Constants
int_2 = '2.0'
int_3 = '3.0'
int_4 = '4.0'
int_5 = '5.0'
label_perc2 = '% severity levels 1-2'
label_perc3 = '% severity level 3'
label_perc4 = '% severity level 4'
label_perc5 = '% severity level 5'
label_tot2 = '# severity levels 1-2'
label_tot3 = '# severity level 3'
label_tot4 = '# severity level 4'
label_tot5 = '# severity level 5'
label_perc_tot = '% Tot PiN (severity levels 3-5)'
label_tot = '# Tot PiN (severity levels 3-5)'
label_admin_severity = 'Area severity'
label_tot_population = 'TotN'

int_acc = 'access'
int_agg= 'aggravating circumstances'
int_lc = 'learning condition'
int_penv = 'protected environment'
int_out = 'Not in need'
label_perc_acc = '% Access'
label_perc_agg= '% Aggravating circumstances'
label_perc_lc = '% Learning conditions'
label_perc_penv = '% Protected environment'
label_perc_out = '% Not in need'
label_tot_acc = '# Access'
label_tot_agg= '# Aggravating circumstances'
label_tot_lc = '# Learning conditions'
label_tot_penv = '# Protected environment'
label_tot_out = '# Not in need'
label_dimension_perc_tot = '% Tot in PiN Dimensions'
label_dimension_tot = '# Tot in PiN Dimensions'
label_dimension_tot_population = 'TotN'


colors = {
    "light_beige": "#FFF2CC",
    "light_orange": "#F4B183",
    "dark_orange": "#ED7D31",
    "darker_orange": "#C65911",
    "light_blue": "#DDEBF7",
    "light_pink": "#b3b389",
    "light_yellow": "#ffffc5",
    "white": "#FFFFFF",
    "bluepin": "#004bb4",
    'gray': '#e0e0e0',
    'stratagray': '#F0F0F0'
}
# Define the columns to color
color_mapping = {
    label_perc2: colors["light_beige"],
    label_tot2: colors["light_beige"],
    label_perc3: colors["light_orange"],
    label_tot3: colors["light_orange"],
    label_perc4: colors["dark_orange"],
    label_tot4: colors["dark_orange"],
    label_perc5: colors["darker_orange"],
    label_tot5: colors["darker_orange"],
    label_perc_tot: colors["light_blue"],
    label_admin_severity: colors["light_blue"],
    label_tot: colors["light_blue"]
}
# Define the colors
colors_dimension = {
    "light_beige": "#ebecc7",
    "light_orange": "#c7ebec",
    "dark_orange": "#c7d9ec",
    "darker_orange": "#c7ecdb",
    'darker2_orange':'#b3d3d4',
    "light_blue": "#DDEBF7",
    "light_pink": "#b3b389",
    "light_yellow": "#ffffc5",
    "white": "#FFFFFF",
    "bluepin": "#004bb4",
    'gray': '#e0e0e0',
    'stratagray': '#F0F0F0'
}

# Define the columns to color
color_mapping_dimension = {
    label_perc2: colors_dimension["light_beige"],
    label_tot2: colors_dimension["light_beige"],
    label_perc3: colors_dimension["dark_orange"],
    label_tot3: colors_dimension["dark_orange"],
    label_perc4: colors_dimension["darker_orange"],
    label_tot4: colors_dimension["darker_orange"],
    label_perc5: colors_dimension["darker2_orange"],
    label_tot5: colors_dimension["darker2_orange"],
    label_perc_tot: colors_dimension["light_blue"],
    label_admin_severity: colors_dimension["light_blue"],
    label_tot: colors_dimension["light_blue"]
}

color_dim ={
    'no_need' : "#a3a380",
    'access' : "#d6ce93",
    'learning' : "#efebce",        
    'envir' : "#d8a48f",
    'aggra' : "#bb8588",    
}
group_gender ={
    'Girls (5-17 y.o.)' : "#3c1642",
    "Boys (5-17 y.o.)": "#1dd3b0"
}
color_school =[
    "#ffbc42",
    "#d81159",
     "#8f2d56",        
     "#218380",
     "#73d2de",    
]

def format_number(num):
    """Convert a large number to a shorter format with K or M."""
    if num >= 1_000_000:
        return f"{num / 1_000_000:.1f}M"
    elif num >= 1_000:
        return f"{num / 1_000:.0f}K"
    else:
        return f"{num:.0f}"

def set_row_height(row, height):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))  # Height in twips (1/20 of a point)
    trHeight.set(qn('w:hRule'), 'exact')  # Exact height
    trPr.append(trHeight)
def set_col_width(column, width):
    for cell in column.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width))
        tcW.set(qn('w:type'), 'dxa')  # 'dxa' means it's in twips, 1/20 of a point
        tcPr.append(tcW)

def create_dimension_table(doc, label, perc_acc, num_acc, perc_lc, num_lc, perc_env, num_env, perc_agg, num_agg, total_in_need, font = 12):
    # Create the table with 5 rows and 3 columns
    table_dimension = doc.add_table(rows=5, cols=3)

    # Set preferred widths for columns
    if font == 12: 
        set_col_width(table_dimension.columns[0],7500)  
        set_col_width(table_dimension.columns[1], 1000)  # Roughly 1 inch
        set_col_width(table_dimension.columns[2], 1000)  # Roughly 1 inch
        set_row_height(table_dimension.rows[0], 350)  
        set_row_height(table_dimension.rows[1], 350) 
        set_row_height(table_dimension.rows[2], 350)  
        set_row_height(table_dimension.rows[3], 350)  
        set_row_height(table_dimension.rows[4], 350)  

    if font != 12: 
        set_col_width(table_dimension.columns[0],6000)  
        set_col_width(table_dimension.columns[1], 700)  # Roughly 1 inch
        set_col_width(table_dimension.columns[2], 700)  # Roughly 1 inch   
        set_row_height(table_dimension.rows[0], 250)  
        set_row_height(table_dimension.rows[1], 250) 
        set_row_height(table_dimension.rows[2], 250)  
        set_row_height(table_dimension.rows[3], 250)  
        set_row_height(table_dimension.rows[4], 250)  

    # Merge the first row across all columns
    cell_title = table_dimension.cell(0, 0)
    cell_title.text = f'Among the {format_number(total_in_need)} {label} children in need:'

    # Left align text in the merged cell
    cell_title.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = cell_title.paragraphs[0].runs[0]
    run_title.font.bold = True
    run_title.font.size = Pt(font)
    run_title.font.name = "Calibri"

    # Apply background color to the merged cell
    cell_title._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["light_blue"])))

    # Add headers to the second row
    table_dimension.cell(0, 1).text = '%'
    table_dimension.cell(0, 2).text = '#'

    # Prepare content to be added to the table
    content = [
        ["Lack access to school", f"{perc_acc}%", f"{format_number(num_acc)}"],
        ["Face unacceptable learning conditions", f"{perc_lc}%", f"{format_number(num_lc)}"],
        ["Live in an unprotected environment in relationship with education", f"{perc_env}%", f"{format_number(num_env)}"],
        ["Face aggravating circumstances preventing their access to education", f"{perc_agg}%", f"{format_number(num_agg)}"],
    ]

    # Add content to the remaining rows
    for i, row in enumerate(content, start=1):
        table_dimension.cell(i, 0).text = row[0]
        table_dimension.cell(i, 1).text = row[1]
        table_dimension.cell(i, 2).text = row[2]

    # Set font and alignment for all cells
    for row in table_dimension.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.size = Pt(font)
                run.font.name = "Calibri"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Specifically set the alignment for the first column (column 0) to the left
    for cell in table_dimension.columns[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align text to the left

    return table_dimension

####################################################################################################################################################################
####################################################################################################################################################################
def create_snapshot_PiN(country_label, final_overview_df, final_overview_df_OCHA,final_overview_dimension_df, final_overview_dimension_df_in_need):
    country_name = country_label.split('__')[0]  # Extract the part before the "__"

    not_pop_group_columns_overview = ['TOTAL (5-17 y.o.)']
    not_pop_group_columns = ['TOTAL (5-17 y.o.)',
        "Girls (5-17 y.o.)", "Boys (5-17 y.o.)", "Female", "Male", "ECE (5 y.o.)", 
        "Primary school", "Intermediate school-level", "Secondary school", "Children with disability"
    ]
    dimension_strata = [ "Girls (5-17 y.o.)", "Boys (5-17 y.o.)", "Female", "Male", "ECE (5 y.o.)", 
        "Primary school", "Intermediate school-level", "Secondary school"
    ]
    
    dimension_gender_strata = [ "Girls (5-17 y.o.)", "Boys (5-17 y.o.)"
    ]
    
    school_cycle_dimension_strata = [ "ECE (5 y.o.)", 
        "Primary school", "Intermediate school-level", "Secondary school"]

    ## reading values for different sessions
    # Retrieve data for the total row
    row_tot = final_overview_df.loc[final_overview_df['Strata'] == 'TOTAL (5-17 y.o.)']
    total_population = row_tot[label_tot_population].values[0]
    tot_perc_in_need = row_tot[label_perc_tot].values[0]
    tot_in_need = row_tot[label_tot].values[0]
    tot_perc_2 = row_tot[label_perc2].values[0]
    tot_perc_3 = row_tot[label_perc3].values[0]
    tot_perc_4 = row_tot[label_perc4].values[0]
    tot_perc_5 = row_tot[label_perc5].values[0]
    tot_num_2 = row_tot[label_tot2].values[0]
    tot_num_3 = row_tot[label_tot3].values[0]
    tot_num_4 = row_tot[label_tot4].values[0]
    tot_num_5 = row_tot[label_tot5].values[0]

    # Retrieve data for various disaggregation
    row_girl = final_overview_df_OCHA.loc[final_overview_df_OCHA['Strata'] == 'Girls (5-17 y.o.)']
    row_boy= final_overview_df_OCHA.loc[final_overview_df_OCHA['Strata'] == 'Boys (5-17 y.o.)']
    row_ece = final_overview_df_OCHA.loc[final_overview_df_OCHA['Strata'] == 'ECE (5 y.o.)']
    row_disability = final_overview_df_OCHA.loc[final_overview_df_OCHA['Strata'] == 'Children with disability']
    tot_in_need_girl = row_girl[label_tot].values[0]
    tot_in_need_boy = row_boy[label_tot].values[0]
    tot_in_need_ece = row_ece[label_tot].values[0]
    tot_in_need_disability = row_disability[label_tot].values[0]

    row_dimension_in_need = final_overview_dimension_df_in_need.loc[final_overview_dimension_df_in_need['Strata'] == 'TOTAL (5-17 y.o.)']
    tot_dimension_in_need = row_dimension_in_need[label_tot_population].values[0]
    perc_acc_in_need = row_dimension_in_need[label_perc_acc].values[0]
    num_acc_in_need = row_dimension_in_need[label_tot_acc].values[0]
    perc_agg_in_need = row_dimension_in_need[label_perc_agg].values[0]
    num_agg_in_need = row_dimension_in_need[label_tot_agg].values[0]
    perc_env_in_need = row_dimension_in_need[label_perc_penv].values[0]
    num_env_in_need = row_dimension_in_need[label_tot_penv].values[0]
    perc_lc_in_need = row_dimension_in_need[label_perc_lc].values[0]
    num_lc_in_need = row_dimension_in_need[label_tot_lc].values[0]

    ###################### plot pop_group

    # Data collection for the bar chart
    groups = []
    severity_3_groups = []
    severity_4_groups = []
    severity_5_groups = []
    percentage_tot_values = []  # Collecting percentage totals for labels
    
    # Iterate over the DataFrame rows to create the content
    for _, row_pop in final_overview_df.iterrows():
        population_group = row_pop['Population group'].upper()  # Convert to uppercase
        strata = row_pop['Strata']
        percentage_3_pop = row_pop[label_perc3]
        percentage_4_pop = row_pop[label_perc4]
        percentage_5_pop = row_pop[label_perc5]
        percentage_tot = row_pop[label_perc_tot]  # Get the total percentage
        
        if strata not in not_pop_group_columns_overview:
            # Append values to the lists for plotting
            groups.append(population_group)
            severity_3_groups.append(percentage_3_pop)
            severity_4_groups.append(percentage_4_pop)
            severity_5_groups.append(percentage_5_pop)
            percentage_tot_values.append(percentage_tot)

    ind = np.arange(len(groups))
    width = 0.5
        
        # Updated code with larger legend text and adjusted layout
    fig, ax = plt.subplots(figsize=(10, 7))  # Adjust the figure size as needed

    # Stacking the bars
    p3 = ax.bar(ind, severity_3_groups, width, color=colors_dimension['dark_orange'], label='Severity 3')
    p4 = ax.bar(ind, severity_4_groups, width, bottom=severity_3_groups, color=colors_dimension['darker_orange'], label='Severity 4')
    p5 = ax.bar(ind, severity_5_groups, width, bottom=np.array(severity_3_groups) + np.array(severity_4_groups), color=colors_dimension['darker2_orange'], label='Severity 5')

    # Adjust the ylim to provide more space at the top
    ax.set_ylim(0, max(np.array(severity_3_groups) + np.array(severity_4_groups) + np.array(severity_5_groups)) + 8)

    # Setting the y-axis label with a larger font size
    ax.set_ylabel('Percentage of children in need', fontsize=16)
    # Setting the x-tick labels with a larger font size
    ax.set_xticks(ind)
    ax.set_xticklabels(groups, fontsize=18)
    # You can also increase the size of the y-tick labels
    ax.tick_params(axis='y', labelsize=12)

    for i, percentage_tot in enumerate(percentage_tot_values):
        ax.text(ind[i], severity_3_groups[i] + severity_4_groups[i] + severity_5_groups[i] + 1,  # 1 is for padding above the bar
                f'{percentage_tot:.1f}%', ha='center', va='bottom', fontsize=15)

    # Adding grid lines
    ax.yaxis.grid(True, linestyle='--', which='major', color='gray', alpha=0.7)

    # Customize the legend with larger text
    legend_labels = [
        'Severity 3:\nOoS children NOT enduring aggravating \ncircumstances or accessing education \nin poor learning conditions',
        'Severity 4:\nOoS and enduring aggravating \ncircumstances or accessing with \nserious protection concerns',
        'Severity 5:\nOoS and enduring exceptional aggravating \ncircumstances or accessing with \nlife-threatening concerns'
    ]

    # Customizing the text in the legend to make "Severity" bold by splitting and recombining the text
    formatted_legend_labels = [
        f'$\mathbf{{Severity\ 3}}$:\nOoS children NOT enduring aggravating \ncircumstances or accessing education \nin poor learning conditions',
        f'$\mathbf{{Severity\ 4}}$:\nOoS and enduring aggravating \ncircumstances or accessing with \nserious protection concerns',
        f'$\mathbf{{Severity\ 5}}$:\nOoS and enduring exceptional aggravating \ncircumstances or accessing with \nlife-threatening concerns'
    ]


    legend = ax.legend([p3, p4, p5], formatted_legend_labels, loc='center left', bbox_to_anchor=(1, 0.5),
                    fontsize=14, title="Severity definition", title_fontsize='14', frameon=True, ncol=1,
                    prop={'size': 14}, handlelength=3, handleheight=4)

    # Adjust the plot and legend layout to give the plot more space
    plt.subplots_adjust(right=0.70)  # Adjust this value to change the ratio between the plot and legend space

    # Save the bar chart
    bar_chart_path = "stacked_bar_chart_large_legend.jpeg"
    plt.savefig(bar_chart_path, format='jpeg', bbox_inches='tight', dpi=300)  # Higher DPI for better quality
    plt.close(fig)



######################
    # Lists to store data
      # Lists to store data
    groups_dimension = []
    no_need_groups = []
    acc_groups = []
    lc_groups = []
    env_groups = []
    agg_groups = []

    # Iterate over the DataFrame rows to collect the content
    for _, row_pop_dim in final_overview_dimension_df.iterrows():
        strata = row_pop_dim['Strata']
        perc_no_need = row_pop_dim[label_perc_out]
        perc_acc = row_pop_dim[label_perc_acc]
        perc_agg = row_pop_dim[label_perc_agg]
        perc_env = row_pop_dim[label_perc_penv]
        perc_lc = row_pop_dim[label_perc_lc]

        # Ensure that we're only adding strata that are relevant
        if strata in dimension_strata:     
            groups_dimension.append(strata)
            no_need_groups.append(perc_no_need)
            acc_groups.append(perc_acc)
            lc_groups.append(perc_lc)
            env_groups.append(perc_env)
            agg_groups.append(perc_agg)

    groups_dimension.reverse()
    no_need_groups.reverse()
    acc_groups.reverse()
    lc_groups.reverse()
    env_groups.reverse()
    agg_groups.reverse()

    # Convert to numpy arrays for easier stacking
    ind_dim = np.arange(len(groups_dimension))
    width = 0.8  # Increase the bar height

    fig_dim, ax_dim = plt.subplots(figsize=(18, 18))  # Increase the height of the figure

    # Stacking the bars, starting with the 'Not in need' group first
    p_no_need = ax_dim.barh(ind_dim, no_need_groups, height=width, color=color_dim['no_need'], label='Not in need')
    p_acc = ax_dim.barh(ind_dim, acc_groups, height=width, left=np.array(no_need_groups), color=color_dim["access"], label='Lack of access')
    p_lc = ax_dim.barh(ind_dim, lc_groups, height=width, left=np.array(no_need_groups) + np.array(acc_groups), color=color_dim["learning"], label='Poor learning conditions')
    p_env = ax_dim.barh(ind_dim, env_groups, height=width, left=np.array(no_need_groups) + np.array(acc_groups) + np.array(lc_groups), color=color_dim["envir"], label='Unprotected environment')
    p_agg = ax_dim.barh(ind_dim, agg_groups, height=width, left=np.array(no_need_groups) + np.array(acc_groups) + np.array(lc_groups) + np.array(env_groups), color=color_dim["aggra"], label='Aggravating circumstances')

    # Adding labels on the bars
    for i in range(len(groups_dimension)):
        if no_need_groups[i] >= 1:
            fontsize = 13 if no_need_groups[i] < 2 else 18
            ax_dim.text(no_need_groups[i] / 2, ind_dim[i], f'{no_need_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)
        if acc_groups[i] >= 1:
            fontsize = 13 if acc_groups[i] < 2 else 18
            ax_dim.text(no_need_groups[i] + acc_groups[i] / 2, ind_dim[i], f'{acc_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)
        if lc_groups[i] >= 1:
            fontsize = 13 if lc_groups[i] < 2 else 18
            ax_dim.text(no_need_groups[i] + acc_groups[i] + lc_groups[i] / 2, ind_dim[i], f'{lc_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)
        if env_groups[i] >= 1:
            fontsize = 13 if env_groups[i] < 2 else 18
            ax_dim.text(no_need_groups[i] + acc_groups[i] + lc_groups[i] + env_groups[i] / 2, ind_dim[i], f'{env_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)
        if agg_groups[i] >= 1:
            fontsize = 13 if agg_groups[i] < 2 else 18
            ax_dim.text(no_need_groups[i] + acc_groups[i] + lc_groups[i] + env_groups[i] + agg_groups[i] / 2, ind_dim[i], f'{agg_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)

    # Remove the x-axis and its gridlines
    ax_dim.get_xaxis().set_visible(False)

    # Keep the y-axis labels but remove the y-axis line
    ax_dim.spines['left'].set_visible(False)  # Hide the left spine (y-axis line)
    ax_dim.spines['right'].set_visible(False)  # Hide the right spine
    ax_dim.spines['top'].set_visible(False)  # Hide the top spine
    ax_dim.spines['bottom'].set_visible(False)  # Hide the bottom spine (x-axis line)

    # Customize the y-axis labels (group names)
    ax_dim.set_yticks(ind_dim)
    ax_dim.set_yticklabels(groups_dimension, fontsize=18, rotation=45, ha='right')  # Increase the font size and rotate labels

    # Customize the legend with larger text
    legend_labels = [
        'Not in need of education (severity 1-2)',
        'OoS, lacking access to education',
        'In-school children, studying in unacceptable basic learning conditions',
        'In-school children, evolving in a non-protective education environment',
        "OoS, facing individual aggravating circumstances"
    ]

    # Move the legend below the plot
    legend = ax_dim.legend([p_no_need, p_acc, p_lc, p_env, p_agg], legend_labels, loc='upper center', bbox_to_anchor=(0.5, -0.15),
                        fontsize=14, title="", title_fontsize='16', frameon=True, ncol=2,
                        prop={'size': 14}, handlelength=3, handleheight=4)
    # Add the small note below the bars and above the legend
    plt.text(0.5, -0.1, "* if the percentage of the need is <1%, the value is not reported in the graph",
            ha='center', va='center', fontsize=12, transform=ax_dim.transAxes)
    # Adjust the plot and legend layout
    plt.subplots_adjust(left=0.05, right=0.95, top=0.95, bottom=0.2)  # Reduce margins

    # Save the bar chart
    bar_chart_path_dim = "horizontal_stacked_bar_chart_high_bars.jpeg"
    plt.savefig(bar_chart_path_dim, format='jpeg', bbox_inches='tight', dpi=300)  # Higher DPI for better quality
    plt.close(fig_dim)





    markertaile = 17

    # Assuming `dimension_gender_strata` and `group_gender` are defined somewhere in your code
    # Lists to store data
    groups_dimension_gender = []
    no_need_groups_gender = []
    acc_groups_gender = []
    lc_groups_gender = []
    env_groups_gender = []
    agg_groups_gender = []

    # Iterate over the DataFrame rows to collect the content
    for _, row_pop_dim_gender in final_overview_dimension_df.iterrows():
        strata_gender = row_pop_dim_gender['Strata']
        perc_no_need_gender = row_pop_dim_gender[label_perc_out]
        perc_acc_gender = row_pop_dim_gender[label_perc_acc]
        perc_agg_gender = row_pop_dim_gender[label_perc_agg]
        perc_env_gender = row_pop_dim_gender[label_perc_penv]
        perc_lc_gender = row_pop_dim_gender[label_perc_lc]

        # Ensure that we're only adding strata that are relevant
        if strata_gender in dimension_gender_strata:     
            groups_dimension_gender.append(strata_gender)
            no_need_groups_gender.append(perc_no_need_gender)
            acc_groups_gender.append(perc_acc_gender)
            lc_groups_gender.append(perc_lc_gender)
            env_groups_gender.append(perc_env_gender)
            agg_groups_gender.append(perc_agg_gender)

    # Reverse the lists to maintain the correct order
    #groups_dimension_gender.reverse()
    #no_need_groups_gender.reverse()
    #acc_groups_gender.reverse()
    #lc_groups_gender.reverse()
    #env_groups_gender.reverse()
    #agg_groups_gender.reverse()

    # Define need labels and group colors
    need_labels = [
        'Not in need: severity 1-2',
        'Lack access to school, OoS',
        'Face unacceptable \nlearning conditions',
        'Live in an unprotected \nenvironment in relationship \nwith education',
        "Face aggravating \ncircumstances preventing \ntheir access to education, OoS"
    ]
    # Create the plot
    fig_dim_gender, ax_dim_gender = plt.subplots(figsize=(14, 10))  # Adjust size as needed

    # Plot each group's data for all needs
    
    for i, group in enumerate(groups_dimension_gender):

        ax_dim_gender.plot([no_need_groups_gender[i]], [need_labels[0]], 'o', color=group_gender[group], markersize=markertaile, label=group)
        ax_dim_gender.plot([acc_groups_gender[i]], [need_labels[1]], 'o', color=group_gender[group], markersize=markertaile)
        ax_dim_gender.plot([lc_groups_gender[i]], [need_labels[2]], 'o', color=group_gender[group], markersize=markertaile)
        ax_dim_gender.plot([env_groups_gender[i]], [need_labels[3]], 'o', color=group_gender[group], markersize=markertaile)
        ax_dim_gender.plot([agg_groups_gender[i]], [need_labels[4]], 'o', color=group_gender[group], markersize=markertaile)


    # Customize the plot
    ax_dim_gender.set_xlabel('Percentage of children affected by:', fontsize=18)
    ax_dim_gender.set_ylabel('')
    ax_dim_gender.set_title('')
    ax_dim_gender.set_xlim(0, max(no_need_groups_gender + acc_groups_gender + lc_groups_gender + env_groups_gender + agg_groups_gender) + 5)
    ax_dim_gender.set_yticks(np.arange(len(need_labels)))
    ax_dim_gender.set_yticklabels(need_labels, fontsize=20)
    ax_dim_gender.tick_params(axis='x', labelsize=16)

    # Add a grid for better readability
    ax_dim_gender.xaxis.grid(True, linestyle='--', which='major', color='gray', alpha=0.7)

    # Add a legend for groups
    ax_dim_gender.legend(
        title="Gender", 
        title_fontsize=18, 
        loc='upper right', 
        bbox_to_anchor=(1, 1), 
        fontsize=18,
        borderaxespad=0.
    )

    # Adjust plot layout for compactness
    plt.subplots_adjust(left=0.15, right=0.85, top=0.9, bottom=0.1)

    # Save the plot
    bar_chart_path_dim_gender = "gender_need.jpeg"
    plt.savefig(bar_chart_path_dim_gender, format='jpeg', bbox_inches='tight', dpi=300)  # Save with high quality
    plt.close(fig_dim_gender)











    # Lists to store data
    groups_dimension_school = []
    no_need_groups_school = []
    acc_groups_school = []
    lc_groups_school = []
    env_groups_school = []
    agg_groups_school = []

    # Iterate over the DataFrame rows to collect the content
    for _, row_pop_dim_school in final_overview_dimension_df.iterrows():
        strata_school = row_pop_dim_school['Strata']
        perc_no_need_school = row_pop_dim_school[label_perc_out]
        perc_acc_school = row_pop_dim_school[label_perc_acc]
        perc_agg_school = row_pop_dim_school[label_perc_agg]
        perc_env_school = row_pop_dim_school[label_perc_penv]
        perc_lc_school = row_pop_dim_school[label_perc_lc]

        # Ensure that we're only adding strata that are relevant
        if strata_school in school_cycle_dimension_strata:     
            groups_dimension_school.append(strata_school)
            no_need_groups_school.append(perc_no_need_school)
            acc_groups_school.append(perc_acc_school)
            lc_groups_school.append(perc_lc_school)
            env_groups_school.append(perc_env_school)
            agg_groups_school.append(perc_agg_school)

    # Reverse the lists to maintain the correct order
    #groups_dimension_school.reverse()
    #no_need_groups_school.reverse()
    #acc_groups_school.reverse()
    #lc_groups_school.reverse()
    #env_groups_school.reverse()
    #agg_groups_school.reverse()

    # Define need labels and group colors

    # Create the plot
    fig_dim_school, ax_dim_school = plt.subplots(figsize=(14, 10))  # Adjust size as needed

    # Plot each group's data for all needs
    
    for i, group in enumerate(groups_dimension_school):
        ax_dim_school.plot([no_need_groups_school[i]], [need_labels[0]], 'o', color=color_school[i], markersize=markertaile, label=group)
        ax_dim_school.plot([acc_groups_school[i]], [need_labels[1]], 'o', color=color_school[i], markersize=markertaile)
        ax_dim_school.plot([lc_groups_school[i]], [need_labels[2]], 'o', color=color_school[i], markersize=markertaile)
        ax_dim_school.plot([env_groups_school[i]], [need_labels[3]], 'o', color=color_school[i], markersize=markertaile)
        ax_dim_school.plot([agg_groups_school[i]], [need_labels[4]], 'o', color=color_school[i], markersize=markertaile)


    # Customize the plot
    ax_dim_school.set_xlabel('Percentage of children affected by:', fontsize=18)
    ax_dim_school.set_ylabel('')
    ax_dim_school.set_title('')
    ax_dim_school.set_xlim(0, max(no_need_groups_school + acc_groups_school + lc_groups_school + env_groups_school + agg_groups_school) + 5)
    ax_dim_school.set_yticks(np.arange(len(need_labels)))
    ax_dim_school.set_yticklabels(need_labels, fontsize=20)
    ax_dim_school.tick_params(axis='x', labelsize=16)

    # Add a grid for better readability
    ax_dim_school.xaxis.grid(True, linestyle='--', which='major', color='gray', alpha=0.7)

    # Add a legend for groups
    ax_dim_school.legend(
        title="School-aged", 
        title_fontsize=18, 
        loc='upper right', 
        bbox_to_anchor=(1, 1), 
        fontsize=18,
        borderaxespad=0.
    )


    # Adjust plot layout for compactness
    plt.subplots_adjust(left=0.15, right=0.85, top=0.9, bottom=0.1)

    # Save the plot
    bar_chart_path_dim_school = "school_need.jpeg"
    plt.savefig(bar_chart_path_dim_school, format='jpeg', bbox_inches='tight', dpi=300)  # Save with high quality
    plt.close(fig_dim_school)





    ###############################################################################################################
    # Initialize the Word document
    doc = docx.Document()

    # Main title style
    title = doc.add_paragraph(f'{country_name} – PiN insight')
    title_run = title.runs[0]
    title_run.font.size = Pt(24)  # Make the title larger
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title.alignment = 1  # Center the title

    # Add some spacing after the title
    doc.add_paragraph()

    # Session 1: Children in Need (5-17 y.o.)
    section1 = doc.add_heading('Children in Need (5-17 y.o.)', level=2)
    section1_run = section1.runs[0]
    section1_run.font.size = Pt(20)  # Customize the section header size
    section1_run.bold = True
    section1_run.font.name = 'Calibri'
    section1.alignment = 0  # Left align
    #doc.add_paragraph()

    paragraph_tot = doc.add_paragraph()

    # Add the first part with the percentage in need, making it bold
    run_perc_in_need = paragraph_tot.add_run(f"{tot_perc_in_need}% (i.e. {format_number(tot_in_need)})")
    run_perc_in_need.font.size = Pt(15)
    run_perc_in_need.font.color.rgb = RGBColor(86,150,151)
    run_perc_in_need.bold = True

    # Add the remaining text of the first sentence (not bold)
    run_text = paragraph_tot.add_run(" of school-aged children are in need. Among them:")
    run_text.font.size = Pt(15)

    # Create a bullet point for each disaggregation with indentation
    bullet_point_girls_boys = doc.add_paragraph(style='List Bullet')
    bullet_point_format = bullet_point_girls_boys.paragraph_format
    bullet_point_format.left_indent = Inches(1)  # Adjust this value for the desired indentation
    bullet_text = bullet_point_girls_boys.add_run(f"{format_number(tot_in_need_girl)} are girls and {format_number(tot_in_need_boy)} are boys;")
    bullet_text.font.size = Pt(12)
    bullet_text.font.name = 'Calibri'

    # Add separate bullet points for other groups with indentation
    bullet_point_ece = doc.add_paragraph(style='List Bullet')
    bullet_point_format = bullet_point_ece.paragraph_format
    bullet_point_format.left_indent = Inches(1)
    bullet_text = bullet_point_ece.add_run(f"{format_number(tot_in_need_ece)} are 5 years old (ECE);")
    bullet_text.font.size = Pt(12)
    bullet_text.font.name = 'Calibri'

    bullet_point_disability = doc.add_paragraph(style='List Bullet')
    bullet_point_format = bullet_point_disability.paragraph_format
    bullet_point_format.left_indent = Inches(1)
    bullet_text = bullet_point_disability.add_run(f"{format_number(tot_in_need_disability)} are children with disability.")
    bullet_text.font.size = Pt(12)
    bullet_text.font.name = 'Calibri'

    pop_group_intro = doc.add_paragraph("School-aged children in need, according to population groups:")
    pop_group_intro.font.size = Pt(12)

    table_pop_group = doc.add_table(rows=1, cols=2)

    # Set the table headers
    table_pop_group.cell(0, 0).text = 'Population groups'
    table_pop_group.cell(0, 1).text = 'Number of school-aged children in need of education'

    # Apply style to the headers
    for cell in table_pop_group.rows[0].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Calibri'

    # Iterate over the DataFrame rows to fill the table for each population group
    for _, row_pop in final_overview_df_OCHA.iterrows():
        total_population_in_need = row_pop[label_tot]
        strata = row_pop['Strata']
        
        if strata not in not_pop_group_columns:
            # Remove the substring '(5-17 y.o.)' and convert to uppercase
            strata_cleaned = strata.replace('(5-17 y.o.)', '').strip().upper()
            
            # Add a new row for each population group
            row_cells = table_pop_group.add_row().cells
            row_cells[0].text = strata_cleaned
            row_cells[1].text = format_number(total_population_in_need)

            # Center align the text and apply font style
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(12)
                run.font.name = 'Calibri'

    # Optionally add a total row at the end
    total_row = table_pop_group.add_row().cells
    total_row[0].text = 'TOTAL'
    total_row[1].text = format_number(tot_in_need)

    # Style the total row
    for cell in total_row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
        run.bold = True  # Bold the total row






    # Add a subsection before the table
    section_severity_distribution = doc.add_heading('Severity Distribution', level=3)
    section_severity_distribution = section_severity_distribution.runs[0]
    section_severity_distribution.font.size = Pt(18)  # Customize the section header size
    section_severity_distribution.font.name = 'Calibri'
    section_severity_distribution.alignment = 0  # Left align
    doc.add_paragraph()

    ################################################### table

    # Add a table with 3 rows and 9 columns
    table = doc.add_table(rows=3, cols=9)

    # Set column widths (you can adjust these as needed)
    table.columns[0].width = Inches(1.5)
    for i in range(1, 9):
        table.columns[i].width = Inches(0.7)

    # Merge cells for the first row
    cell_all_population = table.cell(0, 0)
    #cell_all_population.merge(table.cell(1, 0))
    cell_severity_1_2 = table.cell(0, 1)
    cell_severity_1_2.merge(table.cell(0, 2))
    cell_severity_3 = table.cell(0, 3)
    cell_severity_3.merge(table.cell(0, 4))
    cell_severity_4 = table.cell(0, 5)
    cell_severity_4.merge(table.cell(0, 6))
    cell_severity_5 = table.cell(0, 7)
    cell_severity_5.merge(table.cell(0, 8))

    # Add text to the first row
    cell_all_population.text = 'ToT #'
    cell_severity_1_2.text = 'Severity 1-2'
    cell_severity_3.text = 'Severity 3'
    cell_severity_4.text = 'Severity 4'
    cell_severity_5.text = 'Severity 5'

    # Center the text horizontally and vertically, and remove bold formatting
    for cell in [cell_all_population, cell_severity_1_2, cell_severity_3, cell_severity_4, cell_severity_5]:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w'))))
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(11)
        run.font.bold = False  # Remove bold formatting

    # Add text to the second row
    table.cell(1, 1).text = '%'
    table.cell(1, 2).text = '#'
    table.cell(1, 3).text = '%'
    table.cell(1, 4).text = '#'
    table.cell(1, 5).text = '%'
    table.cell(1, 6).text = '#'
    table.cell(1, 7).text = '%'
    table.cell(1, 8).text = '#'

    # Add numbers to the third row by converting them to strings
    table.cell(2, 0).text = str(format_number(total_population))
    table.cell(2, 1).text = str(tot_perc_2)
    table.cell(2, 2).text = str(format_number(tot_num_2))
    table.cell(2, 3).text = str(tot_perc_3)
    table.cell(2, 4).text = str(format_number(tot_num_3))
    table.cell(2, 5).text = str(tot_perc_4)
    table.cell(2, 6).text = str(format_number(tot_num_4))
    table.cell(2, 7).text = str(tot_perc_5)
    table.cell(2, 8).text = str(format_number(tot_num_5))

    # Apply background colors to the appropriate cells
    for i in range(0, 3):
        for j in [1, 2]:
            table.cell(i, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["light_beige"])))

    for i in range(0, 3):
        for j in [3, 4]:
            table.cell(i, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["dark_orange"])))

    for i in range(0, 3):
        for j in [5, 6]:
            table.cell(i, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["darker_orange"])))

    for i in range(0, 3):
        for j in [7, 8]:
            table.cell(i, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["darker2_orange"])))

    # Iterate over all rows in the table
    for row in table.rows:
        # Iterate over all cells in the current row
        for cell in row.cells:
            # Align the text horizontally in the center
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Align the text vertically in the center
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w'))))


    # Set different heights for the rows
    set_row_height(table.rows[0], 400)  
    set_row_height(table.rows[1], 300) 
    set_row_height(table.rows[2], 300)  


    ##############################################################
    section_population_groups = doc.add_heading('PiN by Severity and by Population Groups', level=3)
    section_population_groups = section_population_groups.runs[0]
    section_population_groups.font.size = Pt(18)  # Customize the section header size
    section_population_groups.font.name = 'Calibri'
    section_population_groups.alignment = 0  # Left align


    # Insert the bar chart after the corresponding section
    doc.add_picture(bar_chart_path, width=Inches(6))  # Adjust the width as needed


    ##############################################################
    section_needs_in_need = doc.add_heading('Insights into the types of needs', level=3)
    section_needs_in_need = section_needs_in_need.runs[0]
    section_needs_in_need.font.size = Pt(18)  # Customize the section header size
    section_needs_in_need.font.name = 'Calibri'
    section_needs_in_need.alignment = 0  # Left align

    doc.add_paragraph("Considering only the children in need, it is possible to explore the needs driving the severity.")

    ## table need
    create_dimension_table(
        doc=doc,
        label="all",
        perc_acc=perc_acc_in_need, num_acc=num_acc_in_need,
        perc_lc= perc_lc_in_need, num_lc=num_lc_in_need,
        perc_env= perc_env_in_need, num_env=num_env_in_need,
        perc_agg=perc_agg, num_agg=num_agg_in_need,
        total_in_need=tot_dimension_in_need,
        font = 12
    )
    doc.add_paragraph("")
    ## table need by pop_group
    for _, row_pop in final_overview_dimension_df_in_need.iterrows():
        population_group = row_pop['Population group'].upper()  # Convert to uppercase
        strata = row_pop['Strata']
        perc_acc = row_pop[label_perc_acc]
        num_acc = row_pop[label_tot_acc]
        perc_agg = row_pop[label_perc_agg]
        num_agg = row_pop[label_tot_agg]
        perc_env = row_pop[label_perc_penv]
        num_env = row_pop[label_tot_penv]
        perc_lc = row_pop[label_perc_lc]
        num_lc = row_pop[label_tot_lc]
        tot_dimension_in_need = row_pop[label_tot_population]
        if strata not in not_pop_group_columns_overview:     
            create_dimension_table(
                doc=doc,
                label=population_group,
                perc_acc=perc_acc, num_acc=num_acc,
                perc_lc= perc_lc, num_lc=num_lc,
                perc_env= perc_env, num_env=num_env,
                perc_agg=perc_agg, num_agg=num_agg,
                total_in_need=tot_dimension_in_need,
                font = 9
            )
            doc.add_paragraph("")

    ##############################################################
    section_needs_msna = doc.add_heading("Children's profile", level=3)
    section_needs_msna = section_needs_msna.runs[0]
    section_needs_msna.font.size = Pt(18)  # Customize the section header size
    section_needs_msna.font.name = 'Calibri'
    section_needs_msna.alignment = 0  # Left align

    doc.add_paragraph("A comprehensive understanding of the profile of children in need can be achieved by examining the four primary dimensions of need: Access, Learning Conditions, Protected Environment, and Aggravating Circumstances.")




    # Insert the bar chart after the corresponding section
    doc.add_picture(bar_chart_path_dim, width=Inches(6))  # Adjust the width as needed
    doc.add_picture(bar_chart_path_dim_gender, width=Inches(6))  # Adjust the width as needed

    doc.add_picture(bar_chart_path_dim_school, width=Inches(6))  # Adjust the width as needed
















    # Save the Word document to a BytesIO object
    doc_output = BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)

    # Return the BytesIO object containing the Word document
    return doc_output

