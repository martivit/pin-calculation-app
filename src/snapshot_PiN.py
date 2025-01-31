import pandas as pd
from fuzzywuzzy import process
import numpy as np
import datetime
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Border, Side, Font, Alignment
from openpyxl.cell.cell import MergedCell
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
import matplotlib.pyplot as plt
from io import BytesIO
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib.text import Text
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib as mpl



# Set the global font to Calibri
mpl.rcParams['font.family'] = 'Calibri'
mpl.rcParams['font.size'] = 12  # Set default font size if needed


# Constants
int_2 = '2.0'
int_3 = '3.0'
int_4 = '4.0'
int_5 = '5.0'
label_perc2 = '% severity levels 1-2'
label_perc3 = '% severity level 3'
label_perc4 = '% severity level 4'
label_perc5 = '% severity level 5'
label_tot2 = '# severity levels 1-2'
label_tot3 = '# severity level 3'
label_tot4 = '# severity level 4'
label_tot5 = '# severity level 5'
label_perc_tot = '% Tot PiN (severity levels 3-5)'
label_tot = '# Tot PiN (severity levels 3-5)'
label_admin_severity = 'Area severity'
label_tot_population = 'TotN'

int_acc = 'access'
int_agg= 'aggravating circumstances'
int_lc = 'learning condition'
int_penv = 'protected environment'
int_out = 'Not in need'
label_perc_acc = '% Access'
label_perc_agg= '% Aggravating circumstances'
label_perc_lc = '% Learning conditions'
label_perc_penv = '% Protected environment'
label_perc_out = '% Not in need'
label_tot_acc = '# Access'
label_tot_agg= '# Aggravating circumstances'
label_tot_lc = '# Learning conditions'
label_tot_penv = '# Protected environment'
label_tot_out = '# Not in need'
label_dimension_perc_tot = '% Tot in PiN Dimensions'
label_dimension_tot = '# Tot in PiN Dimensions'
label_dimension_tot_population = 'TotN'


colors = {
    "light_beige": "#FFF2CC",
    "light_orange": "#F4B183",
    "dark_orange": "#ED7D31",
    "darker_orange": "#C65911",
    "light_blue": "#DDEBF7",
    "light_pink": "#b3b389",
    "light_yellow": "#ffffc5",
    "white": "#FFFFFF",
    "bluepin": "#004bb4",
    'gray': '#e0e0e0',
    'stratagray': '#F0F0F0'
}
# Define the columns to color
color_mapping = {
    label_perc2: colors["light_beige"],
    label_tot2: colors["light_beige"],
    label_perc3: colors["light_orange"],
    label_tot3: colors["light_orange"],
    label_perc4: colors["dark_orange"],
    label_tot4: colors["dark_orange"],
    label_perc5: colors["darker_orange"],
    label_tot5: colors["darker_orange"],
    label_perc_tot: colors["light_blue"],
    label_admin_severity: colors["light_blue"],
    label_tot: colors["light_blue"]
}
# Define the colors
colors_dimension = {
    "light_beige": "#ebecc7",
    "light_orange": "#c7ebec",
    "dark_orange": "#c7d9ec",
    "darker_orange": "#c7ecdb",
    'darker2_orange':'#b3d3d4',
    "light_blue": "#DDEBF7",
    "light_pink": "#b3b389",
    "light_yellow": "#ffffc5",
    "white": "#FFFFFF",
    "bluepin": "#004bb4",
    'gray': '#e0e0e0',
    'stratagray': '#F0F0F0'
}

# Define the columns to color
color_mapping_dimension = {
    label_perc2: colors_dimension["light_beige"],
    label_tot2: colors_dimension["light_beige"],
    label_perc3: colors_dimension["dark_orange"],
    label_tot3: colors_dimension["dark_orange"],
    label_perc4: colors_dimension["darker_orange"],
    label_tot4: colors_dimension["darker_orange"],
    label_perc5: colors_dimension["darker2_orange"],
    label_tot5: colors_dimension["darker2_orange"],
    label_perc_tot: colors_dimension["light_blue"],
    label_admin_severity: colors_dimension["light_blue"],
    label_tot: colors_dimension["light_blue"]
}

color_dim ={
    'no_need' : "#a3a380",
    'access' : "#d6ce93",
    'learning' : "#efebce",        
    'envir' : "#d8a48f",
    'aggra' : "#bb8588",    
}
group_gender ={
    'Girls (5-17 y.o.)' : "#3c1642",
    "Boys (5-17 y.o.)": "#1dd3b0"
}
group_gender_afg ={
    'Girls (6-17 y.o.)' : "#3c1642",
    "Boys (6-17 y.o.)": "#1dd3b0"
}



color_school =[
    "#ffbc42",
    "#d81159",
     "#8f2d56",        
     "#218380",
     "#73d2de",    
]

def format_number(num):
    """Convert a large number to a shorter format with K or M."""
    if num >= 1_000_000:
        return f"{num / 1_000_000:.1f}M"
    elif num >= 1_000:
        return f"{num / 1_000:.1f}K"
    else:
        return f"{num:.1f}"

def set_row_height(row, height):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))  # Height in twips (1/20 of a point)
    trHeight.set(qn('w:hRule'), 'exact')  # Exact height
    trPr.append(trHeight)
def set_col_width(column, width):
    for cell in column.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width))
        tcW.set(qn('w:type'), 'dxa')  # 'dxa' means it's in twips, 1/20 of a point
        tcPr.append(tcW)

def set_top_bottom_border(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    # Top border
    top_border = OxmlElement('w:top')
    top_border.set('w:val', 'single')
    top_border.set('w:sz', '6')
    top_border.set('w:space', '0')
    top_border.set('w:color', '000000')
    tcBorders.append(top_border)

    # Bottom border
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set('w:val', 'single')
    bottom_border.set('w:sz', '6')
    bottom_border.set('w:space', '0')
    bottom_border.set('w:color', '000000')
    tcBorders.append(bottom_border)

    tcPr.append(tcBorders)




# Call this function after creating the document title
image_path = 'pics/pin_table.jpg'  # Path to your uploaded image
textbox_text = (
    "Education PiN definition\n"
    "The definition agreed globally and as part of the JIAF of a child in need is the following:\n"
    "Number of school-aged children and youths in the areas affected by crisis who do not have access to protective education and acceptable learning conditions, which can negatively impact:\n"
)
bullet_points = [
    "(i) their physical and psychosocial wellbeing",
    "(ii) cognitive development",
    "(iii) ability to meet their future needs."
]
additional_text = (
    "\nThe severity of the situation for Education is computed based on 4 dimensions (ALPA):\n"
    "• Access to education\n"
    "• Learning conditions\n"
    "• Protection level in or on the way to school\n"
    "• Aggravating circumstances (at individual level)\n\n"
    "The table below recaps some examples of indicators falling under each of these dimensions and the main data sources "
    "usually available in countries where clusters are activated. The list is not exhaustive, and other indicators deemed "
    "appropriate at country level can be used."
)


## -------------------------------------------------------------------------------------------------------------------------------------------------
def create_dimension_table(doc, label, perc_acc, num_acc, perc_lc, num_lc, perc_env, num_env, perc_agg, num_agg, total_in_need, font = 12):
    # Create the table with 5 rows and 3 columns
    table_dimension = doc.add_table(rows=5, cols=3)

    # Set preferred widths for columns
    if font == 12: 
        set_col_width(table_dimension.columns[0],7500)  
        set_col_width(table_dimension.columns[1], 1000)  # Roughly 1 inch
        set_col_width(table_dimension.columns[2], 1000)  # Roughly 1 inch
        set_row_height(table_dimension.rows[0], 350)  
        set_row_height(table_dimension.rows[1], 350) 
        set_row_height(table_dimension.rows[2], 350)  
        set_row_height(table_dimension.rows[3], 350)  
        set_row_height(table_dimension.rows[4], 350)  

    if font != 12: 
        set_col_width(table_dimension.columns[0],6000)  
        set_col_width(table_dimension.columns[1], 700)  # Roughly 1 inch
        set_col_width(table_dimension.columns[2], 700)  # Roughly 1 inch   
        set_row_height(table_dimension.rows[0], 250)  
        set_row_height(table_dimension.rows[1], 250) 
        set_row_height(table_dimension.rows[2], 250)  
        set_row_height(table_dimension.rows[3], 250)  
        set_row_height(table_dimension.rows[4], 250)  

    # Merge the first row across all columns
    cell_title = table_dimension.cell(0, 0)
    cell_title.text = f'Among {label} children in need ({format_number(total_in_need)}):'

    # Left align text in the merged cell
    cell_title.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = cell_title.paragraphs[0].runs[0]
    run_title.font.bold = True
    run_title.font.size = Pt(font)
    run_title.font.name = "Calibri"

    # Apply background color to the merged cell
    cell_title._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["light_blue"])))

    # Add headers to the second row
    table_dimension.cell(0, 1).text = '%'
    table_dimension.cell(0, 2).text = '#'

    # Prepare content to be added to the table
    content = [
        ["Lack access to school", f"{perc_acc}%", f"{format_number(num_acc)}"],
        ["Study in unacceptable basic learning conditions", f"{perc_lc}%", f"{format_number(num_lc)}"],
        ["Evolve in a non-protective education environment", f"{perc_env}%", f"{format_number(num_env)}"],
        ["Face aggravating circumstances preventing their access to education", f"{perc_agg}%", f"{format_number(num_agg)}"],
    ]
    severity_colors = ['#f4a582', '#fddbc7', '#fee8c8', '#fff7ec']
    percentages = [perc_acc, perc_lc, perc_env, perc_agg]
    sorted_percentages = sorted(range(len(percentages)), key=lambda k: percentages[k], reverse=True)

    # Add content to the remaining rows
    for i, row in enumerate(content, start=1):
        table_dimension.cell(i, 0).text = row[0]
        table_dimension.cell(i, 1).text = row[1]
        table_dimension.cell(i, 2).text = row[2]
        # Apply color coding to the entire row
         # Get the color based on the sorted order of percentages
        color_index = sorted_percentages.index(i-1)
        color = severity_colors[color_index]

        # Apply color coding to the entire row
        for j in [1, 2]:  # Columns 1 and 2
            table_dimension.cell(i, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)))


    # Set font and alignment for all cells
    for row in table_dimension.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.size = Pt(font)
                run.font.name = "Calibri"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Specifically set the alignment for the first column (column 0) to the left
    for cell in table_dimension.columns[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align text to the left

    table_dimension.style = 'Table Grid'

    return table_dimension

## -------------------------------------------------------------------------------------------------------------------------------------------------
# Filter the data for each section
def filter_data(df, strata_list):
    filtered_data = df[df['Strata'].isin(strata_list)]
    return filtered_data
## -------------------------------------------------------------------------------------------------------------------------------------------------
# Function to plot each section
def plot_snapshot(ax, df, title, color_dim, fixed_height):
    groups_dimension = []
    no_need_groups = []
    acc_groups = []
    lc_groups = []
    env_groups = []
    agg_groups = []

    # Iterate over the DataFrame rows to collect the content
    for _, row_pop_dim in df.iterrows():
        strata = row_pop_dim['Strata']
        perc_no_need = row_pop_dim[label_perc_out]
        perc_acc = row_pop_dim[label_perc_acc]
        perc_agg = row_pop_dim[label_perc_agg]
        perc_env = row_pop_dim[label_perc_penv]
        perc_lc = row_pop_dim[label_perc_lc]

        groups_dimension.append(strata)
        no_need_groups.append(perc_no_need)
        acc_groups.append(perc_acc)
        lc_groups.append(perc_lc)
        env_groups.append(perc_env)
        agg_groups.append(perc_agg)

    # Reverse lists
    groups_dimension.reverse()
    no_need_groups.reverse()
    acc_groups.reverse()
    lc_groups.reverse()
    env_groups.reverse()
    agg_groups.reverse()

    ind_dim = np.arange(len(groups_dimension))
    width = fixed_height  # Set the fixed bar height

    # Stacking the bars, starting with the 'Not in need' group first
    p_no_need = ax.barh(ind_dim, no_need_groups, height=width, color=color_dim['no_need'], label='Not in need')
    p_acc = ax.barh(ind_dim, acc_groups, height=width, left=np.array(no_need_groups), color=color_dim["access"], label='Lack of access')
    p_lc = ax.barh(ind_dim, lc_groups, height=width, left=np.array(no_need_groups) + np.array(acc_groups), color=color_dim["learning"], label='Poor learning conditions')
    p_env = ax.barh(ind_dim, env_groups, height=width, left=np.array(no_need_groups) + np.array(acc_groups) + np.array(lc_groups), color=color_dim["envir"], label='Unprotected environment')
    p_agg = ax.barh(ind_dim, agg_groups, height=width, left=np.array(no_need_groups) + np.array(acc_groups) + np.array(lc_groups) + np.array(env_groups), color=color_dim["aggra"], label='Aggravating circumstances')

    for i in range(len(groups_dimension)):
        total_offset = 0  # Initialize total offset for stacking

        if no_need_groups[i] >= 1:
            fontsize = 13 if no_need_groups[i] < 2 else 18
            ax.text(total_offset + no_need_groups[i] / 2, ind_dim[i], f'{no_need_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)
            total_offset += no_need_groups[i]  # Update offset

        if acc_groups[i] >= 1:
            fontsize = 13 if acc_groups[i] < 2 else 18
            ax.text(total_offset + acc_groups[i] / 2, ind_dim[i], f'{acc_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)
            total_offset += acc_groups[i]  # Update offset

        if lc_groups[i] >= 1:
            fontsize = 13 if lc_groups[i] < 2 else 18
            ax.text(total_offset + lc_groups[i] / 2, ind_dim[i], f'{lc_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)
            total_offset += lc_groups[i]  # Update offset

        if env_groups[i] >= 1:
            fontsize = 13 if env_groups[i] < 2 else 18
            ax.text(total_offset + env_groups[i] / 2, ind_dim[i], f'{env_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)
            total_offset += env_groups[i]  # Update offset

        if agg_groups[i] >= 1:
            fontsize = 13 if agg_groups[i] < 2 else 18
            ax.text(total_offset + agg_groups[i] / 2, ind_dim[i], f'{agg_groups[i]:.1f}%', ha='center', va='center', fontsize=fontsize, color='black', rotation=45)
            total_offset += agg_groups[i]  # Update offset

    # Customize the y-axis labels (group names)
    ax.set_yticks(ind_dim)
    ax.set_yticklabels(groups_dimension, fontsize=18, rotation=45, ha='right')  # Increase the font size and rotate labels
    ax.set_title(title, fontsize=20)

    ax.get_xaxis().set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    ax.spines['bottom'].set_visible(False)

    # Return the bar handles for the legend
    return p_no_need, p_acc, p_lc, p_env, p_agg
####################################################################################################################################################################
####################################################################################################################################################################
def create_snapshot_PiN(country_label, final_overview_df, final_overview_df_OCHA,final_overview_dimension_df, final_overview_dimension_df_in_need, selected_language= 'English'):
    country_name = country_label.split('__')[0]  # Extract the part before the "__"

    print(country_name)

    tot_5_17_label = 'TOTAL (5-17 y.o.)'
    girl_5_17_label = 'Girls (5-17 y.o.)'
    boy_5_17_label = 'Boys (5-17 y.o.)'
    ece_5yo_label = 'ECE (5 y.o.)'
    if country_name == 'Afghanistan':
        tot_5_17_label = 'TOTAL (6-17 y.o.)'
        girl_5_17_label = 'Girls (6-17 y.o.)'
        boy_5_17_label = 'Boys (6-17 y.o.)'
        ece_5yo_label = 'ECE (6 y.o.)'

    not_pop_group_columns_overview = [tot_5_17_label]
    not_pop_group_columns = [tot_5_17_label,
        girl_5_17_label,boy_5_17_label, "Female", "Male", ece_5yo_label, 
        "Primary school", "Intermediate school-level", "Secondary school", "Children with disability"
    ]
    dimension_strata = [ girl_5_17_label, boy_5_17_label, "Female", "Male", ece_5yo_label, 
        "Primary school", "Intermediate school-level", "Secondary school"
    ]
    
    dimension_gender_strata = [ girl_5_17_label, boy_5_17_label]
    
    school_cycle_dimension_strata = [ ece_5yo_label, 
        "Primary school", "Intermediate school-level", "Secondary school"]
    school_cycle_dimension_strata_wo_ece = ["Primary school", "Intermediate school-level", "Secondary school"]
    ece_strata = [ece_5yo_label]

    ## reading values for different sessions
    # Retrieve data for the total row
    print(final_overview_df)
    row_tot = final_overview_df.loc[final_overview_df['Strata'] == tot_5_17_label]
    print(tot_5_17_label)
    print(row_tot)
    total_population = row_tot[label_tot_population].values[0]
    tot_perc_in_need = row_tot[label_perc_tot].values[0]
    tot_in_need = row_tot[label_tot].values[0]
    tot_perc_2 = row_tot[label_perc2].values[0]
    tot_perc_3 = row_tot[label_perc3].values[0]
    tot_perc_4 = row_tot[label_perc4].values[0]
    tot_perc_5 = row_tot[label_perc5].values[0]
    tot_num_2 = row_tot[label_tot2].values[0]
    tot_num_3 = row_tot[label_tot3].values[0]
    tot_num_4 = row_tot[label_tot4].values[0]
    tot_num_5 = row_tot[label_tot5].values[0]

    # Retrieve data for various disaggregation
    row_girl = final_overview_df_OCHA.loc[final_overview_df_OCHA['Strata'] == girl_5_17_label]
    row_boy= final_overview_df_OCHA.loc[final_overview_df_OCHA['Strata'] == boy_5_17_label]
    row_ece = final_overview_df_OCHA.loc[final_overview_df_OCHA['Strata'] == ece_5yo_label]
    row_disability = final_overview_df_OCHA.loc[final_overview_df_OCHA['Strata'] == 'Children with disability']
    tot_in_need_girl = row_girl[label_tot].values[0]
    tot_in_need_boy = row_boy[label_tot].values[0]
    tot_in_need_ece = row_ece[label_tot].values[0]
    tot_in_need_disability = row_disability[label_tot].values[0]

    row_dimension_in_need = final_overview_dimension_df_in_need.loc[final_overview_dimension_df_in_need['Strata'] == tot_5_17_label]
    tot_dimension_in_need = row_dimension_in_need[label_tot_population].values[0]
    perc_acc_in_need = row_dimension_in_need[label_perc_acc].values[0]
    num_acc_in_need = row_dimension_in_need[label_tot_acc].values[0]
    perc_agg_in_need = row_dimension_in_need[label_perc_agg].values[0]
    num_agg_in_need = row_dimension_in_need[label_tot_agg].values[0]
    perc_env_in_need = row_dimension_in_need[label_perc_penv].values[0]
    num_env_in_need = row_dimension_in_need[label_tot_penv].values[0]
    perc_lc_in_need = row_dimension_in_need[label_perc_lc].values[0]
    num_lc_in_need = row_dimension_in_need[label_tot_lc].values[0]

    ###################### plot pop_group

    # Data collection for the bar chart
    groups = []
    severity_3_groups = []
    severity_4_groups = []
    severity_5_groups = []
    percentage_tot_values = []  # Collecting percentage totals for labels
    text_data = []  

    
    # Iterate over the DataFrame rows to create the content
    for _, row_pop in final_overview_df.iterrows():
        population_group = row_pop['Population group'].upper()  # Convert to uppercase
        strata = row_pop['Strata']
        percentage_3_pop = row_pop[label_perc3]
        percentage_4_pop = row_pop[label_perc4]
        percentage_5_pop = row_pop[label_perc5]
        percentage_tot = row_pop[label_perc_tot]  # Get the total percentage
        
        if strata not in not_pop_group_columns_overview:
            # Append values to the lists for plotting
            groups.append(population_group)
            severity_3_groups.append(percentage_3_pop)
            severity_4_groups.append(percentage_4_pop)
            severity_5_groups.append(percentage_5_pop)
            percentage_tot_values.append(percentage_tot)
            text_line = f'{population_group}: {percentage_3_pop:.1f}% (severity 3), {percentage_4_pop:.1f}% (severity 4), {percentage_5_pop:.1f}% (severity 5)'
            text_data.append(text_line)  # Append the formatted string to the list
    text_str = '\n'.join(text_data)  # Join the strings with newlines


    ind = np.arange(len(groups))
    width = 0.5
        
        # Updated code with larger legend text and adjusted layout
    fig, ax = plt.subplots(figsize=(10, 7))  # Adjust the figure size as needed

    # Stacking the bars
    p3 = ax.bar(ind, severity_3_groups, width, color=colors_dimension['dark_orange'], label='Severity 3')
    p4 = ax.bar(ind, severity_4_groups, width, bottom=severity_3_groups, color=colors_dimension['darker_orange'], label='Severity 4')
    p5 = ax.bar(ind, severity_5_groups, width, bottom=np.array(severity_3_groups) + np.array(severity_4_groups), color=colors_dimension['darker2_orange'], label='Severity 5')

    # Adjust the ylim to provide more space at the top
    ax.set_ylim(0, max(np.array(severity_3_groups) + np.array(severity_4_groups) + np.array(severity_5_groups)) + 8)

    # Setting the y-axis label with a larger font size
    ax.set_ylabel('Percentage of children in need', fontsize=16)
    # Setting the x-tick labels with a larger font size
    ax.set_xticks(ind)
    ax.set_xticklabels(groups, fontsize=18)
    # You can also increase the size of the y-tick labels
    ax.tick_params(axis='y', labelsize=12)

    for i, percentage_tot in enumerate(percentage_tot_values):
        ax.text(ind[i], severity_3_groups[i] + severity_4_groups[i] + severity_5_groups[i] + 1,  # 1 is for padding above the bar
                f'{percentage_tot:.1f}%', ha='center', va='bottom', fontsize=16)

    # Adding grid lines
    ax.yaxis.grid(True, linestyle='--', which='major', color='gray', alpha=0.7)

    # Customizing the text in the legend to make "Severity" bold by splitting and recombining the text
    formatted_legend_labels = [
        f'$\mathbf{{Severity\ level\ 3}}$:\nOoS children who do NOT endure aggravating \ncircumstances OR in-school children \nstudying in unacceptable basic learning conditions',
        f'$\mathbf{{Severity\ level\ 4}}$:\nOoS children who endure aggravating \ncircumstances OR in-school children evolving in a \nnon-protective education environment',
        f'$\mathbf{{Severity\ level\ 5}}$:\nOoS children who endure exceptionally aggravating \ncircumstances OR in-school children evolving \nin a life-threating education environment'
    ]


    legend = ax.legend([p3, p4, p5], formatted_legend_labels, loc='center left', bbox_to_anchor=(1, 0.5),
                    fontsize=14, title="Severity definitions", title_fontsize='14', frameon=True, ncol=1,
                    prop={'size': 14}, handlelength=3, handleheight=4)
    
    ax.text(1.05, 0.15, text_str, transform=ax.transAxes, fontsize=13, verticalalignment='top', 
        bbox=dict(facecolor='white', alpha=0.5), wrap=True)

    # Adjust the plot and legend layout to give the plot more space
    plt.subplots_adjust(right=0.70)  # Adjust this value to change the ratio between the plot and legend space

    # Save the bar chart
    bar_chart_path = "stacked_bar_chart_large_legend.jpeg"
    plt.savefig(bar_chart_path, format='jpeg', bbox_inches='tight', dpi=300)  # Higher DPI for better quality
    plt.close(fig)



######################

    # Plotting three sections in one figure
# Calculate the number of bars in each subplot

    # Filter the data for each section
    df_gender = filter_data(final_overview_dimension_df, dimension_gender_strata)
    df_ece = filter_data(final_overview_dimension_df, ece_strata)
    df_school_cycles = filter_data(final_overview_dimension_df, school_cycle_dimension_strata_wo_ece)

    num_bars_gender = len(df_gender)
    num_bars_ece = len(df_ece)
    num_bars_school_cycles = len(df_school_cycles)

    # Proportionally allocate heights based on the number of bars in each subplot
    total_bars = num_bars_gender + num_bars_ece + num_bars_school_cycles
    height_ratio_gender = num_bars_gender / total_bars
    height_ratio_ece = num_bars_ece / total_bars
    height_ratio_school_cycles = num_bars_school_cycles / total_bars

    # Plotting three sections in one figure with proportional heights for subplots
    fig_subsection, axs_subsection = plt.subplots(
        3, 1, figsize=(15, 18), 
        gridspec_kw={'height_ratios': [height_ratio_gender, height_ratio_ece, height_ratio_school_cycles]}
    )

    fixed_bar_height = 0.8  # You can adjust this value as needed
    label_snapshot_5yo = "Snapshot of ECE (5 y.o.)"
    if country_name == 'Afghanistan':
        label_snapshot_5yo = "Snapshot of ECE (6 y.o.)"

    # Plot each section
    p_no_need, p_acc, p_lc, p_env, p_agg = plot_snapshot(axs_subsection[0], df_gender, "Snapshot of gender", color_dim, fixed_bar_height)
    plot_snapshot(axs_subsection[1], df_ece, label_snapshot_5yo, color_dim, fixed_bar_height)
    plot_snapshot(axs_subsection[2], df_school_cycles, "Snapshot of school cycles", color_dim, fixed_bar_height)


    # Add the shared legend at the bottom
    legend_labels = [
        'Not in need of education (severity 1-2)',
        'OoS, lacking access to education',
        'In-school children, studying in unacceptable basic learning conditions',
        'In-school children, evolving in a non-protective education environment',
        "OoS, facing individual aggravating circumstances"
    ]

    fig_subsection.legend([p_no_need, p_acc, p_lc, p_env, p_agg], labels=legend_labels, loc='upper center', bbox_to_anchor=(0.5, -0.05),
            fontsize=14, ncol=2, prop={'size': 14}, handlelength=3, handleheight=4)
    
    plt.text(0.5, -0.1, "* if the percentage of the need is <1%, the value is not reported in the graph",
         ha='center', va='center', fontsize=14, transform=axs_subsection[2].transAxes)
    # Adjust the layout
    plt.subplots_adjust(left=0.05, right=0.95, top=0.95, bottom=0.15)  # Adjust margins for space

    # Save or display the plot
    bar_chart_path_subsection = "snapshot_grouped_horizontal_bar_chart.jpeg"
    plt.savefig(bar_chart_path_subsection, format='jpeg', bbox_inches='tight', dpi=300)
    plt.close(fig_subsection)



    markertaile = 17

    # Assuming `dimension_gender_strata` and `group_gender` are defined somewhere in your code
    # Lists to store data
    groups_dimension_gender = []
    no_need_groups_gender = []
    acc_groups_gender = []
    lc_groups_gender = []
    env_groups_gender = []
    agg_groups_gender = []

    # Iterate over the DataFrame rows to collect the content
    for _, row_pop_dim_gender in final_overview_dimension_df.iterrows():
        strata_gender = row_pop_dim_gender['Strata']
        perc_no_need_gender = row_pop_dim_gender[label_perc_out]
        perc_acc_gender = row_pop_dim_gender[label_perc_acc]
        perc_agg_gender = row_pop_dim_gender[label_perc_agg]
        perc_env_gender = row_pop_dim_gender[label_perc_penv]
        perc_lc_gender = row_pop_dim_gender[label_perc_lc]

        # Ensure that we're only adding strata that are relevant
        if strata_gender in dimension_gender_strata:     
            groups_dimension_gender.append(strata_gender)
            no_need_groups_gender.append(perc_no_need_gender)
            acc_groups_gender.append(perc_acc_gender)
            lc_groups_gender.append(perc_lc_gender)
            env_groups_gender.append(perc_env_gender)
            agg_groups_gender.append(perc_agg_gender)

    # Reverse the lists to maintain the correct order
    #groups_dimension_gender.reverse()
    #no_need_groups_gender.reverse()
    #acc_groups_gender.reverse()
    #lc_groups_gender.reverse()
    #env_groups_gender.reverse()
    #agg_groups_gender.reverse()

    # Define need labels and group colors
    need_labels = [
        'Not in need of education \n(severity 1-2)',
        'OoS, lacking access to education',
        'In-school children, \nstudying in unacceptable \nbasic learning conditions',
        'In-school children, \nevolving in a non-protective \neducation environment',
        "OoS, facing individual \naggravating circumstances"
    ]
    # Create the plot
    fig_dim_gender, ax_dim_gender = plt.subplots(figsize=(14, 10))  # Adjust size as needed

    # Plot each group's data for all needs
    if country_name != 'Afghanistan':
        for i, group in enumerate(groups_dimension_gender):
            ax_dim_gender.plot([no_need_groups_gender[i]], [need_labels[0]], 'o', color=group_gender[group], markersize=markertaile, label=group)
            ax_dim_gender.plot([acc_groups_gender[i]], [need_labels[1]], 'o', color=group_gender[group], markersize=markertaile)
            ax_dim_gender.plot([lc_groups_gender[i]], [need_labels[2]], 'o', color=group_gender[group], markersize=markertaile)
            ax_dim_gender.plot([env_groups_gender[i]], [need_labels[3]], 'o', color=group_gender[group], markersize=markertaile)
            ax_dim_gender.plot([agg_groups_gender[i]], [need_labels[4]], 'o', color=group_gender[group], markersize=markertaile)
    else:
        for i, group in enumerate(groups_dimension_gender):
            ax_dim_gender.plot([no_need_groups_gender[i]], [need_labels[0]], 'o', color=group_gender_afg[group], markersize=markertaile, label=group)
            ax_dim_gender.plot([acc_groups_gender[i]], [need_labels[1]], 'o', color=group_gender_afg[group], markersize=markertaile)
            ax_dim_gender.plot([lc_groups_gender[i]], [need_labels[2]], 'o', color=group_gender_afg[group], markersize=markertaile)
            ax_dim_gender.plot([env_groups_gender[i]], [need_labels[3]], 'o', color=group_gender_afg[group], markersize=markertaile)
            ax_dim_gender.plot([agg_groups_gender[i]], [need_labels[4]], 'o', color=group_gender_afg[group], markersize=markertaile)



    # Customize the plot
    ax_dim_gender.set_xlabel('Percentage of children affected by:', fontsize=18)
    ax_dim_gender.set_ylabel('')
    ax_dim_gender.set_title('')
    ax_dim_gender.set_xlim(0, max(no_need_groups_gender + acc_groups_gender + lc_groups_gender + env_groups_gender + agg_groups_gender) + 5)
    ax_dim_gender.set_yticks(np.arange(len(need_labels)))
    ax_dim_gender.set_yticklabels(need_labels, fontsize=20)
    ax_dim_gender.tick_params(axis='x', labelsize=16)

    # Add a grid for better readability
    ax_dim_gender.xaxis.grid(True, linestyle='--', which='major', color='gray', alpha=0.7)

    # Add a legend for groups
    ax_dim_gender.legend(
        title="Gender", 
        title_fontsize=18, 
        loc='upper right', 
        bbox_to_anchor=(1, 1), 
        fontsize=18,
        borderaxespad=0.
    )

    # Adjust plot layout for compactness
    plt.subplots_adjust(left=0.15, right=0.85, top=0.9, bottom=0.1)

    # Save the plot
    bar_chart_path_dim_gender = "gender_need.jpeg"
    plt.savefig(bar_chart_path_dim_gender, format='jpeg', bbox_inches='tight', dpi=300)  # Save with high quality
    plt.close(fig_dim_gender)











    # Lists to store data
    groups_dimension_school = []
    no_need_groups_school = []
    acc_groups_school = []
    lc_groups_school = []
    env_groups_school = []
    agg_groups_school = []

    # Iterate over the DataFrame rows to collect the content
    for _, row_pop_dim_school in final_overview_dimension_df.iterrows():
        strata_school = row_pop_dim_school['Strata']
        perc_no_need_school = row_pop_dim_school[label_perc_out]
        perc_acc_school = row_pop_dim_school[label_perc_acc]
        perc_agg_school = row_pop_dim_school[label_perc_agg]
        perc_env_school = row_pop_dim_school[label_perc_penv]
        perc_lc_school = row_pop_dim_school[label_perc_lc]

        # Ensure that we're only adding strata that are relevant
        if strata_school in school_cycle_dimension_strata:     
            groups_dimension_school.append(strata_school)
            no_need_groups_school.append(perc_no_need_school)
            acc_groups_school.append(perc_acc_school)
            lc_groups_school.append(perc_lc_school)
            env_groups_school.append(perc_env_school)
            agg_groups_school.append(perc_agg_school)

    # Reverse the lists to maintain the correct order
    #groups_dimension_school.reverse()
    #no_need_groups_school.reverse()
    #acc_groups_school.reverse()
    #lc_groups_school.reverse()
    #env_groups_school.reverse()
    #agg_groups_school.reverse()

    # Define need labels and group colors

    # Create the plot
    fig_dim_school, ax_dim_school = plt.subplots(figsize=(14, 10))  # Adjust size as needed

    # Plot each group's data for all needs
    
    for i, group in enumerate(groups_dimension_school):
        ax_dim_school.plot([no_need_groups_school[i]], [need_labels[0]], 'o', color=color_school[i], markersize=markertaile, label=group)
        ax_dim_school.plot([acc_groups_school[i]], [need_labels[1]], 'o', color=color_school[i], markersize=markertaile)
        ax_dim_school.plot([lc_groups_school[i]], [need_labels[2]], 'o', color=color_school[i], markersize=markertaile)
        ax_dim_school.plot([env_groups_school[i]], [need_labels[3]], 'o', color=color_school[i], markersize=markertaile)
        ax_dim_school.plot([agg_groups_school[i]], [need_labels[4]], 'o', color=color_school[i], markersize=markertaile)


    # Customize the plot
    ax_dim_school.set_xlabel('Percentage of children affected by:', fontsize=18)
    ax_dim_school.set_ylabel('')
    ax_dim_school.set_title('')
    ax_dim_school.set_xlim(0, max(no_need_groups_school + acc_groups_school + lc_groups_school + env_groups_school + agg_groups_school) + 5)
    ax_dim_school.set_yticks(np.arange(len(need_labels)))
    ax_dim_school.set_yticklabels(need_labels, fontsize=20)
    ax_dim_school.tick_params(axis='x', labelsize=16)

    # Add a grid for better readability
    ax_dim_school.xaxis.grid(True, linestyle='--', which='major', color='gray', alpha=0.7)

    # Add a legend for groups
    ax_dim_school.legend(
        title="School-aged", 
        title_fontsize=18, 
        loc='upper right', 
        bbox_to_anchor=(1, 1), 
        fontsize=18,
        borderaxespad=0.
    )


    # Adjust plot layout for compactness
    plt.subplots_adjust(left=0.15, right=0.85, top=0.9, bottom=0.1)

    # Save the plot
    bar_chart_path_dim_school = "school_need.jpeg"
    plt.savefig(bar_chart_path_dim_school, format='jpeg', bbox_inches='tight', dpi=300)  # Save with high quality
    plt.close(fig_dim_school)





    ###############################################################################################################
    # Initialize the Word document
    doc = docx.Document()

    # Main title style
    title = doc.add_paragraph(f'{country_name} – PiN insight')
    title_run = title.runs[0]
    title_run.font.size = Pt(24)  # Make the title larger
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title.alignment = 1  # Center the title

    # Define the gray background color (hex code)
    gray_background = "e0e0e0"

    # Add a 1x1 table to the document
    table_pin = doc.add_table(rows=1, cols=1)

    # Get the first (and only) cell of the table
    cell_pin = table_pin.cell(0, 0)

    # Insert the main text into the table cell
   # Start a new paragraph for the cell
    paragraph = cell_pin.add_paragraph()

    # Add the bold part for "Education PiN definition"
    run = paragraph.add_run("Education PiN definition")
    run.bold = True
    run.font.size = Pt(10)  # Smaller font size
    run.font.name = 'Calibri'

    # Set paragraph spacing to reduce the extra space
    paragraph.paragraph_format.space_before = Pt(0)  # No space before the paragraph
    paragraph.paragraph_format.space_after = Pt(0)   # No space after the paragraph
    paragraph.paragraph_format.line_spacing = Pt(12)  # Adjust line spacing

    # Add the rest of the text (regular font)
    run = paragraph.add_run("\nThe definition agreed globally and as part of the JIAF of a child in need is the following:\n"
                            "Number of school-aged children and youths in the areas affected by crisis who do not have access to "
                            "protective education and acceptable learning conditions, which can negatively impact:\n")
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    # Add bullet points with reduced spacing
    bullet_points = [
        "(i) their physical and psychosocial wellbeing",
        "(ii) cognitive development",
        "(iii) ability to meet their future needs."
    ]

    for point in bullet_points:
        p = cell_pin.add_paragraph(point, style='List Bullet')
        p.paragraph_format.space_before = Pt(0)  # Reduce space before bullet point
        p.paragraph_format.space_after = Pt(0)   # Reduce space after bullet point
        p.paragraph_format.line_spacing = Pt(12) # Adjust line spacing
        run = p.runs[0]
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # Add additional text after the bullet points with reduced spacing
    additional_text = (
        "\nThe severity of the situation for Education is computed based on 4 dimensions (ALPA):\n"
        "1) Access to education\n"
        "2) Learning conditions\n"
        "3) Protection level in or on the way to school\n"
        "4) Aggravating circumstances (at individual level)\n\n"
        "The table below recaps some examples of indicators falling under each of these dimensions and the main data sources "
        "usually available in countries where clusters are activated. The list is not exhaustive, and other indicators deemed "
        "appropriate at country level can be used."
    )

    additional_paragraph = cell_pin.add_paragraph(additional_text)
    additional_paragraph.paragraph_format.space_before = Pt(0)  # Reduce space before the paragraph
    additional_paragraph.paragraph_format.space_after = Pt(0)   # Reduce space after the paragraph
    additional_paragraph.paragraph_format.line_spacing = Pt(12) # Adjust line spacing

    for run in additional_paragraph.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # Insert the image into the table cell
    paragraph = cell_pin.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6))  # Adjust the width as needed

    # Set the table width to the width of the page
    table_pin.autofit = False
    table_pin.allow_autofit = False
    table_pin.columns[0].width = Inches(6.5)  # Adjust the width based on your page settings

    # Center the table in the document
    table_pin.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Apply a gray background to the cell
    cell_pin._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), gray_background)))

    # Apply font size 10 to the entire table content
    for paragraph in cell_pin.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'


    # Add some spacing after the title
    doc.add_page_break()


    # Session 1: Children in Need (5-17 y.o.)
    if country_name != 'Afghanistan':
        section1 = doc.add_heading(f'Children in Need (5-17 y.o.) in {country_name}', level=2)
    else:
        section1 = doc.add_heading(f'Children in Need (6-17 y.o.) in {country_name}', level=2)
    section1_run = section1.runs[0]
    section1_run.font.size = Pt(20)  # Customize the section header size
    section1_run.bold = True
    section1_run.font.name = 'Calibri'
    section1.alignment = 0  # Left align
    doc.add_paragraph()

    paragraph_tot = doc.add_paragraph()

    # Add the first part with the percentage in need, making it bold
    run_perc_in_need = paragraph_tot.add_run(f"{tot_perc_in_need}% (i.e. {format_number(tot_in_need)})")
    run_perc_in_need.font.size = Pt(16)
    run_perc_in_need.font.name = 'Calibri'
    run_perc_in_need.font.color.rgb = RGBColor(86,150,151)
    run_perc_in_need.bold = True

    # Add the remaining text of the first sentence (not bold)
    run_text = paragraph_tot.add_run(" of school-aged children are in need. Among them:")
    run_text.font.size = Pt(16)
    run_text.font.name = 'Calibri'

    # Create a bullet point for each disaggregation with indentation
    bullet_point_girls_boys = doc.add_paragraph(style='List Bullet')
    bullet_point_format = bullet_point_girls_boys.paragraph_format
    bullet_point_format.left_indent = Inches(1)  # Adjust this value for the desired indentation
    bullet_text = bullet_point_girls_boys.add_run(f"{format_number(tot_in_need_girl)} are girls and {format_number(tot_in_need_boy)} are boys;")
    bullet_text.font.size = Pt(16)
    bullet_text.font.name = 'Calibri'

    # Add separate bullet points for other groups with indentation
    bullet_point_ece = doc.add_paragraph(style='List Bullet')
    bullet_point_format = bullet_point_ece.paragraph_format
    bullet_point_format.left_indent = Inches(1)
    if country_name != 'Afghanistan':
        bullet_text = bullet_point_ece.add_run(f"{format_number(tot_in_need_ece)} are 5 years old (ECE);")
    else:
        bullet_text = bullet_point_ece.add_run(f"{format_number(tot_in_need_ece)} are 6 years old (ECE);")
    bullet_text.font.size = Pt(16)
    bullet_text.font.name = 'Calibri'

        # Add separate bullet points for other groups with indentation
    bullet_point_rest = doc.add_paragraph(style='List Bullet')
    bullet_point_format = bullet_point_rest.paragraph_format
    bullet_point_format.left_indent = Inches(1)
    if country_name != 'Afghanistan':
        bullet_text = bullet_point_rest.add_run(f"{format_number(tot_in_need-tot_in_need_ece)} are between 6 and 17 years old;")
    else:
        bullet_text = bullet_point_rest.add_run(f"{format_number(tot_in_need-tot_in_need_ece)} are between 7 and 17 years old;")

    bullet_text.font.size = Pt(16)
    bullet_text.font.name = 'Calibri'

    bullet_point_disability = doc.add_paragraph(style='List Bullet')
    bullet_point_format = bullet_point_disability.paragraph_format
    bullet_point_format.left_indent = Inches(1)
    bullet_text = bullet_point_disability.add_run(f"{format_number(tot_in_need_disability)} are children with disability.")
    bullet_text.font.size = Pt(16)
    bullet_text.font.name = 'Calibri'

    doc.add_paragraph()

    pop_group_intro = doc.add_paragraph("School-aged children in need, according to population groups:")
    # Apply font and size formatting to the text
    run_pop_group_intro = pop_group_intro.runs[0]
    run_pop_group_intro.font.name = 'Calibri'
    run_pop_group_intro.font.size = Pt(16)

    table_pop_group = doc.add_table(rows=1, cols=2)

    # Set the table headers
    table_pop_group.cell(0, 0).text = 'Population groups'
    table_pop_group.cell(0, 1).text = 'Number of school-aged children in need of education'


    # Apply style to the headers
    for cell in table_pop_group.rows[0].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        #run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Calibri'

    for _, row_pop in final_overview_df_OCHA.iterrows():
        total_population_in_need = row_pop[label_tot]
        strata = row_pop['Strata']
        
        if strata not in not_pop_group_columns:
            # Remove the substring '(5-17 y.o.)' and convert to uppercase
            strata_cleaned = strata.replace('(5-17 y.o.)', '').strip().upper()
            if country_name == 'Afghanistan':
                  strata_cleaned = strata.replace('(6-17 y.o.)', '').strip().upper()


            # Add a new row for each population group
            row_cells = table_pop_group.add_row().cells
            row_cells[0].text = strata_cleaned
            row_cells[1].text = format_number(total_population_in_need)

            # Center align the text and apply font style
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(12)
                run.font.name = 'Calibri'

    # Optionally add a total row at the end
    total_row = table_pop_group.add_row().cells
    total_row[0].text = 'TOTAL'
    total_row[1].text = format_number(tot_in_need)

    # Style the total row
    for cell in total_row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
        #run.bold = True  # Bold the total row

    for row in table_pop_group.rows:
        set_row_height(row, 300) 

    set_row_height(table_pop_group.rows[0], 300)  
    set_col_width(table_pop_group.columns[0],1000)  
    set_col_width(table_pop_group.columns[1], 5000)  # Roughly 1 inch
    table_pop_group.style = 'Table Grid'
    table_pop_group.alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.add_paragraph()

    #pteacher = doc.add_paragraph("Additionally, X of adults (teachers/educational staff) are in need of education assistance.")
    #pteacher_format = pteacher.runs[0]
    #pteacher_format.font.name = 'Calibri'
    #pteacher_format.font.size = Pt(16)

    # Add a subsection before the table
    section_severity_distribution = doc.add_heading('School-aged children severity distribution', level=3)
    section_severity_distribution = section_severity_distribution.runs[0]
    section_severity_distribution.font.size = Pt(18)  # Customize the section header size
    section_severity_distribution.font.name = 'Calibri'
    section_severity_distribution.alignment = 0  # Left align
    doc.add_paragraph()

    ################################################### table

    # Add a table with 3 rows and 9 columns
    table = doc.add_table(rows=4, cols=9)

    # Set column widths (you can adjust these as needed)
    table.columns[0].width = Inches(1.5)
    for i in range(1, 9):
        table.columns[i].width = Inches(0.7)

    # Merge cells for the first row
    cell_pin = table.cell(0, 3)
    cell_pin.merge(table.cell(0, 8))

    cell_all_population = table.cell(1, 0)
    #cell_all_population.merge(table.cell(1, 0))
    cell_severity_1_2 = table.cell(1, 1)
    cell_severity_1_2.merge(table.cell(1, 2))
    cell_severity_3 = table.cell(1, 3)
    cell_severity_3.merge(table.cell(1, 4))
    cell_severity_4 = table.cell(1, 5)
    cell_severity_4.merge(table.cell(1, 6))
    cell_severity_5 = table.cell(1, 7)
    cell_severity_5.merge(table.cell(1, 8))

    # Add text to the first row
    cell_pin.text = 'School-aged children PiN'
    cell_all_population.text = 'ToT #'
    cell_severity_1_2.text = 'Severity levels 1-2'
    cell_severity_3.text = 'Severity level 3'
    cell_severity_4.text = 'Severity level 4'
    cell_severity_5.text = 'Severity level 5'

    # Center the text horizontally and vertically, and remove bold formatting
    for cell in [cell_pin,cell_all_population, cell_severity_1_2, cell_severity_3, cell_severity_4, cell_severity_5]:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w'))))
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.bold = False  # Remove bold formatting

    # Add text to the second row
    table.cell(2, 1).text = '%'
    table.cell(2, 2).text = '#'
    table.cell(2, 3).text = '%'
    table.cell(2, 4).text = '#'
    table.cell(2, 5).text = '%'
    table.cell(2, 6).text = '#'
    table.cell(2, 7).text = '%'
    table.cell(2, 8).text = '#'

    # Add numbers to the third row by converting them to strings
    table.cell(3, 0).text = str(format_number(total_population))
    table.cell(3, 1).text = str(tot_perc_2)
    table.cell(3, 2).text = str(format_number(tot_num_2))
    table.cell(3, 3).text = str(tot_perc_3)
    table.cell(3, 4).text = str(format_number(tot_num_3))
    table.cell(3, 5).text = str(tot_perc_4)
    table.cell(3, 6).text = str(format_number(tot_num_4))
    table.cell(3, 7).text = str(tot_perc_5)
    table.cell(3, 8).text = str(format_number(tot_num_5))

    # Apply background colors to the appropriate cells

    for j in [3, 8]:
        table.cell(0, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors["dark_orange"])))
    for i in range(1, 4):
        for j in [1, 2]:
            table.cell(i, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["light_beige"])))

    for i in range(1, 4):
        for j in [3, 4]:
            table.cell(i, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["dark_orange"])))

    for i in range(1, 4):
        for j in [5, 6]:
            table.cell(i, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["darker_orange"])))

    for i in range(1, 4):
        for j in [7, 8]:
            table.cell(i, j)._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), colors_dimension["darker2_orange"])))

    # Iterate over all rows in the table
    for row in table.rows:
        # Iterate over all cells in the current row
        for cell in row.cells:
            # Align the text horizontally in the center
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Align the text vertically in the center
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w'))))


    # Set different heights for the rows
    set_row_height(table.rows[0], 300)  
    set_row_height(table.rows[1], 600)  
    set_row_height(table.rows[2], 300) 
    set_row_height(table.rows[3], 300)  


    ##############################################################
    section_population_groups = doc.add_heading('Unpacking levels of severity for each population group', level=3)
    section_population_groups = section_population_groups.runs[0]
    section_population_groups.font.size = Pt(18)  # Customize the section header size
    section_population_groups.font.name = 'Calibri'
    section_population_groups.alignment = 0  # Left align


    # Insert the bar chart after the corresponding section
    doc.add_picture(bar_chart_path, width=Inches(6))  # Adjust the width as needed


    ##############################################################
    section_needs_in_need = doc.add_heading('Unpacking children needs', level=3)
    section_needs_in_need = section_needs_in_need.runs[0]
    section_needs_in_need.font.size = Pt(18)  # Customize the section header size
    section_needs_in_need.font.name = 'Calibri'
    section_needs_in_need.alignment = 0  # Left align

    intro_need = doc.add_paragraph("What is driving the severity of children in needs?")
    intro_need_format = intro_need.runs[0]
    intro_need_format.font.name = 'Calibri'
    intro_need_format.font.size = Pt(16)
    # Add the new text with bullet points
    bullet_point_text = (
        "Children in need are categorized based on four key dimensions of the Education People in Need (PiN) framework: "
        "Access to education, Learning conditions, Protection level in or on the way to school, and Aggravating circumstances. "
        "These dimensions determine their severity level:"
    )
    # Add the paragraph for the introductory sentence
    intro_paragraph = doc.add_paragraph(bullet_point_text)
    intro_paragraph_format = intro_paragraph.runs[0]
    intro_paragraph_format.font.name = 'Calibri'
    intro_paragraph_format.font.size = Pt(12)

    # Create bullet points
    bullet_points = [
        "Lack of access to school: This applies to children who fall into severity 3 and are not accessing school at all.",
        "Studying in unacceptable basic learning conditions: These children are assigned to severity 3, but they do have access to school. However, the quality of their learning environment is extremely poor and inadequate.",
        "Evolve in a non-protective education environment: Children assigned to severity 4 or 5 who are attending school but in an environment that does not ensure their safety or protection.",
        "Face aggravating circumstances preventing their access to education: Children assigned to either severity 4 or 5, where external factors significantly affect their ability to attend school, leading to a lack of access to education."
    ]

    # Add each bullet point with formatting
    for point in bullet_points:
        bullet_paragraph = doc.add_paragraph(style='List Bullet')
        bullet_run = bullet_paragraph.add_run(point)
        bullet_run.font.name = 'Calibri'
        bullet_run.font.size = Pt(12)

    ## table need
    create_dimension_table(
        doc=doc,
        label="all",
        perc_acc=perc_acc_in_need, num_acc=num_acc_in_need,
        perc_lc= perc_lc_in_need, num_lc=num_lc_in_need,
        perc_env= perc_env_in_need, num_env=num_env_in_need,
        perc_agg=perc_agg_in_need, num_agg=num_agg_in_need,
        total_in_need=tot_dimension_in_need,
        font = 11
    )
    doc.add_paragraph("")
    ## table need by pop_group
    for _, row_pop in final_overview_dimension_df_in_need.iterrows():
        population_group = row_pop['Population group'].upper()  # Convert to uppercase
        strata = row_pop['Strata']
        perc_acc = row_pop[label_perc_acc]
        num_acc = row_pop[label_tot_acc]
        perc_agg = row_pop[label_perc_agg]
        num_agg = row_pop[label_tot_agg]
        perc_env = row_pop[label_perc_penv]
        num_env = row_pop[label_tot_penv]
        perc_lc = row_pop[label_perc_lc]
        num_lc = row_pop[label_tot_lc]
        tot_dimension_in_need = row_pop[label_tot_population]
        if strata not in not_pop_group_columns_overview:     
            create_dimension_table(
                doc=doc,
                label=population_group,
                perc_acc=perc_acc, num_acc=num_acc,
                perc_lc= perc_lc, num_lc=num_lc,
                perc_env= perc_env, num_env=num_env,
                perc_agg=perc_agg, num_agg=num_agg,
                total_in_need=tot_dimension_in_need,
                font = 9
            )
            doc.add_paragraph("")

    ##############################################################


    section_needs_msna = doc.add_heading("Children's profiles (from MSNA, HH survey)", level=3)
    section_needs_msna = section_needs_msna.runs[0]
    section_needs_msna.font.size = Pt(18)  # Customize the section header size
    section_needs_msna.font.name = 'Calibri'
    section_needs_msna.alignment = 0  # Left align

    intro_dimension = doc.add_paragraph("The graphics below show school-aged children profiles, relying on the 4 dimensions of the education PiN (ALPA, Access to education, Learning conditions, Protection level in or on the way to school and Aggravating circumstances).")
    intro_dimension_format = intro_dimension.runs[0]
    intro_dimension_format.font.name = 'Calibri'
    intro_dimension_format.font.size = Pt(12)
    



    # Insert the bar chart after the corresponding section
    #doc.add_picture(bar_chart_path_dim, width=Inches(6))  # Adjust the width as needed
    doc.add_picture(bar_chart_path_subsection, width=Inches(5))
    doc.add_picture(bar_chart_path_dim_gender, width=Inches(5))  # Adjust the width as needed
    doc.add_picture(bar_chart_path_dim_school, width=Inches(5))  # Adjust the width as needed
















    # Save the Word document to a BytesIO object
    doc_output = BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)

    # Return the BytesIO object containing the Word document
    return doc_output

