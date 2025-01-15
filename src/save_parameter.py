import pandas as pd
from fuzzywuzzy import process
import numpy as np
import datetime
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Border, Side, Font, Alignment
from openpyxl.cell.cell import MergedCell
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
import matplotlib.pyplot as plt
from io import BytesIO
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib.text import Text
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib as mpl
from docx import Document
from datetime import datetime



def generate_parameters(st_session_state):
    """
    Generate the parameters dictionary for PiN calculation.

    Args:
        st_session_state: Streamlit session state object containing relevant variables.

    Returns:
        dict: The generated parameters dictionary.
    """
    parameters = {
        "general_info": {
            "country": st_session_state.get('country'),
            "date_calculation": datetime.now().strftime("%d/%m/%Y %H:%M")  # Today's date with hour and minute
        },
        "msna indicators per PiN dimension": {
            "access": st_session_state.get('access_var'),
            "learning condition": {
                "Education disrupted due to teacher absences": st_session_state.get('teacher_disruption_var'),
                "Education disrupted due to natural hazard": st_session_state.get('natural_hazard_disruption_var'),
            },
            "protected environment": {
                "Education disrupted due to school being used as IDP shelter": st_session_state.get('idp_disruption_var'),
                "Education disrupted due to school being occupied by armed groups": st_session_state.get('armed_disruption_var'),
            },
            "aggravating_circumstances": st_session_state.get('barrier_var'),
        },
        "severity_classification": {
            "severity level 3": {
                "description": "OoS children who do NOT endure aggravating circumstances or in-school children whose education was disrupted due to:",
                "ind1 in-school": st_session_state.get('teacher_disruption_var'),
                "ind2 in-school": st_session_state.get('natural_hazard_disruption_var'),
            },
            "severity level 4": {
                "description": "In-school children whose education disrupted due to or OoS facing the following aggravating circumstances.",
                "ind in-school": st_session_state.get('idp_disruption_var'),
                "aggravating circumstances": st_session_state.get('selected_severity_4_barriers', []),
            },
            "severity level 5": {
                "description": "In-school children whose education disrupted due to or OoS facing the following aggravating circumstances.",
                "ind in-school": st_session_state.get('armed_disruption_var'),
                "aggravating circumstances": st_session_state.get('selected_severity_5_barriers', []),
            },
        },
        "admin_unit": {
            "HNO unit of analysis": st_session_state.get('admin_var'),
            "mismatch admin": st_session_state.get('mismatch_admin', False),
        },
        "school_cycles": {
            "age_ranges": st_session_state.get('vector_cycle'),  # Age groups for educational cycles
        }
    }
    return parameters



def generate_word_document(parameters):
    # Initialize the Word document
    doc = docx.Document()
    doc.add_heading('Parameters Used as Input for the PiN Calculation', level=1)

    # Add General Information
    doc.add_heading('General Information', level=2)
    general_info = parameters["general_info"]
    for key, value in general_info.items():
        doc.add_paragraph(f"{key.replace('_', ' ').capitalize()}: {value}", style='List Bullet')

    # Add MSNA Indicators
    # Add MSNA Indicators
    doc.add_heading('MSNA indicators/variables by dimension', level=2)
    msna_indicators = parameters["msna indicators per PiN dimension"]
    for category, indicators in msna_indicators.items():
        if isinstance(indicators, dict):  # Nested categories
            # Main bullet for the category with bold formatting
            category_paragraph = doc.add_paragraph(style='List Bullet')
            category_run = category_paragraph.add_run(f"{category.replace('_', ' ').capitalize()}:")
            category_run.bold = True
            for description, indicator in indicators.items():
                # Sub-bullets for each indicator
                doc.add_paragraph(f"      {description}: {indicator}", style='List Bullet 2')
        else:
            # Main bullet for simple categories with bold formatting
            category_paragraph = doc.add_paragraph(style='List Bullet')
            category_run = category_paragraph.add_run(f"{category.replace('_', ' ').capitalize()}: {indicators}")
            category_run.bold = True


    # Add Severity Classification
    doc.add_heading('Severity Classification used for this calculation', level=2)
            # Add Severity Classification
    severity_classification = parameters["severity_classification"]
    for level, details in severity_classification.items():
        # Define colors for severity levels
        color_map = {
            "severity level 3": RGBColor(255, 165, 0),  # Light orange
            "severity level 4": RGBColor(255, 140, 0),  # Darker orange
            "severity level 5": RGBColor(255, 69, 0),   # Red-orange
        }

        # Add severity level heading
        severity_paragraph = doc.add_paragraph(style='List Bullet')
        severity_run = severity_paragraph.add_run(f"{level.replace('_', ' ').capitalize()}: ")
        severity_run.bold = True
        if level in color_map:
            severity_run.font.color.rgb = color_map[level]

        # Handle Severity Level 3 with two details
        if level == "severity level 3":
            description = details["description"]
            # Add the description first
            severity_paragraph.add_run(description + " ")
            
            if "ind1 in-school" in details and "ind2 in-school" in details:
                detail_1 = details["ind1 in-school"]
                detail_2 = details["ind2 in-school"]
                
                # Add the first detail in bold
                detail_run1 = severity_paragraph.add_run(detail_1)
                detail_run1.bold = True
                
                # Add " and " between the details
                severity_paragraph.add_run(" and ")
                
                # Add the second detail in bold
                detail_run2 = severity_paragraph.add_run(detail_2)
                detail_run2.bold = True
                
                # Add the period at the end
                severity_paragraph.add_run(".")

        # Handle Severity Levels 4 and 5 with one detail
        elif level in ["severity level 4", "severity level 5"]:
            description = details["description"]
            description_parts = description.split("due to")
            severity_paragraph.add_run(description_parts[0] + "due to ")
            if "ind in-school" in details:
                detail_1 = details["ind in-school"]
                detail_run = severity_paragraph.add_run(detail_1)
                detail_run.bold = True
            if len(description_parts) > 1:
                severity_paragraph.add_run(description_parts[1])

        # Add examples as sub-bullets
        if "aggravating circumstances" in details:
            for example in details["aggravating circumstances"]:
                example_paragraph = doc.add_paragraph(style='List Bullet 2')
                example_paragraph.add_run(f"      {example}")


                    


    # Add Admin Unit
    doc.add_heading('Administrative Unit', level=2)
    admin_unit = parameters["admin_unit"]
    for key, value in admin_unit.items():
        doc.add_paragraph(f"{key.replace('_', ' ').capitalize()}: {value}", style='List Bullet')

    # Add School Cycles
    doc.add_heading('School Cycles', level=2)
    school_cycles = parameters.get("school_cycles", {})
    # Report 'age_ranges' (vector_cycle) as-is
    age_ranges = school_cycles.get("age_ranges", [])
    doc.add_paragraph(f"Age Ranges: {age_ranges}", style='List Bullet')

    # Save the Word document to a BytesIO object
    from io import BytesIO
    doc_output = BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)

    return doc_output
