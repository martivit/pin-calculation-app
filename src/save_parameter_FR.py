import pandas as pd
from fuzzywuzzy import process
import numpy as np
import datetime
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Border, Side, Font, Alignment
from openpyxl.cell.cell import MergedCell
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
import matplotlib.pyplot as plt
from io import BytesIO
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib.text import Text
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib as mpl
from docx import Document
from datetime import datetime

def generate_parameters_FR(st_session_state):
    """
    Générer le dictionnaire des paramètres pour le calcul PiN.

    Args:
        st_session_state: Objet d'état de session Streamlit contenant les variables pertinentes.

    Returns:
        dict: Le dictionnaire des paramètres généré.
    """
    parameters = {
        "informations_generales": {
            "pays": st_session_state.get('country'),
            "date_du_calcul": datetime.now().strftime("%d/%m/%Y %H:%M")  # Date actuelle avec heure et minute
        },
        "indicateurs_msna_par_dimension": {
            "accès": st_session_state.get('access_var'),
            "conditions_d_apprentissage": {
                "Éducation perturbée en raison de l'absence des enseignants": st_session_state.get('teacher_disruption_var'),
                "Éducation perturbée en raison d'un aléa naturel": st_session_state.get('natural_hazard_disruption_var'),
            },
            "environnement_protégé": {
                "Éducation perturbée en raison de l'utilisation de l'école comme abri pour les PDI": st_session_state.get('idp_disruption_var'),
                "Éducation perturbée en raison de l'occupation de l'école par des groupes armés": st_session_state.get('armed_disruption_var'),
            },
            "circonstances_aggravantes": st_session_state.get('barrier_var'),
        },
        "classification_de_sévérité": {
            "niveau_de_sévérité_3": {
                "description": "Enfants hors école qui ne subissent PAS de circonstances aggravantes ou enfants scolarisés dont l'éducation a été perturbée en raison de :",
                "ind1_scolarisés": st_session_state.get('teacher_disruption_var'),
                "ind2_scolarisés": st_session_state.get('natural_hazard_disruption_var'),
            },
            "niveau_de_sévérité_4": {
                "description": "Enfants scolarisés dont l'éducation a été perturbée en raison de (indicateur scolarisé) ou enfants hors école confrontés aux circonstances aggravantes suivantes.",
                "indicateur_scolarisé": st_session_state.get('idp_disruption_var'),
                "circonstances_aggravantes": st_session_state.get('selected_severity_4_barriers', []),
            },
            "niveau_de_sévérité_5": {
                "description": "Enfants scolarisés dont l'éducation a été perturbée en raison de (indicateur scolarisé) ou enfants hors école confrontés aux circonstances aggravantes suivantes.",
                "indicateur_scolarisé": st_session_state.get('armed_disruption_var'),
                "circonstances_aggravantes": st_session_state.get('selected_severity_5_barriers', []),
            },
        },
        "unité_administrative": {
            "unité_d_analyse": st_session_state.get('admin_var'),
            "décalage_admin": st_session_state.get('mismatch_admin', False),
        },
        "cycles_scolaires": {
            "tranches_d_age": st_session_state.get('vector_cycle'),  # Groupes d'âge pour les cycles éducatifs
        }
    }
    return parameters


def generate_word_document_FR(parameters):
    """
    Générer un document Word en français pour les paramètres.

    Args:
        parameters (dict): Dictionnaire des paramètres.

    Returns:
        BytesIO: Document Word généré.
    """
    # Initialiser le document Word
    doc = docx.Document()
    doc.add_heading('Paramètres Utilisés comme Entrée pour le Calcul PiN', level=1)

    # Ajouter les Informations Générales
    doc.add_heading('Informations Générales', level=2)
    general_info = parameters["informations_generales"]
    for key, value in general_info.items():
        doc.add_paragraph(f"{key.replace('_', ' ').capitalize()}: {value}", style='List Bullet')

    # Ajouter les Indicateurs MSNA
    doc.add_heading('Indicateurs/Variables MSNA par Dimension', level=2)
    msna_indicators = parameters["indicateurs_msna_par_dimension"]
    for category, indicators in msna_indicators.items():
        if isinstance(indicators, dict):  # Catégories imbriquées
            category_paragraph = doc.add_paragraph(style='List Bullet')
            category_run = category_paragraph.add_run(f"{category.replace('_', ' ').capitalize()}:")
            category_run.bold = True
            for description, indicator in indicators.items():
                doc.add_paragraph(f"      {description}: {indicator}", style='List Bullet 2')
        else:
            category_paragraph = doc.add_paragraph(style='List Bullet')
            category_run = category_paragraph.add_run(f"{category.replace('_', ' ').capitalize()}: {indicators}")
            category_run.bold = True

    # Ajouter la Classification de Sévérité
    doc.add_heading('Classification de Sévérité Utilisée pour ce Calcul', level=2)
    severity_classification = parameters["classification_de_sévérité"]
    for level, details in severity_classification.items():
        color_map = {
            "niveau_de_sévérité_3": RGBColor(255, 165, 0),
            "niveau_de_sévérité_4": RGBColor(255, 140, 0),
            "niveau_de_sévérité_5": RGBColor(255, 69, 0),
        }

        severity_paragraph = doc.add_paragraph(style='List Bullet')
        severity_run = severity_paragraph.add_run(f"{level.replace('_', ' ').capitalize()}: ")
        severity_run.bold = True
        if level in color_map:
            severity_run.font.color.rgb = color_map[level]

        if level == "niveau_de_sévérité_3":
            description = details["description"]
            severity_paragraph.add_run(description + " ")
            if "ind1_scolarisés" in details and "ind2_scolarisés" in details:
                detail_1 = details["ind1_scolarisés"]
                detail_2 = details["ind2_scolarisés"]
                detail_run1 = severity_paragraph.add_run(detail_1)
                detail_run1.bold = True
                severity_paragraph.add_run(" et ")
                detail_run2 = severity_paragraph.add_run(detail_2)
                detail_run2.bold = True
                severity_paragraph.add_run(".")

        elif level in ["niveau_de_sévérité_4", "niveau_de_sévérité_5"]:
            description = details["description"]
            description = description.replace("en raison de (indicateur scolarisé)", "en raison de")
            description_parts = description.split("en raison de")
            severity_paragraph.add_run(description_parts[0] + "en raison de ")
            if "indicateur_scolarisé" in details:
                detail_1 = details["indicateur_scolarisé"]
                detail_run = severity_paragraph.add_run(detail_1)
                detail_run.bold = True
            if len(description_parts) > 1:
                severity_paragraph.add_run(description_parts[1])

        if "circonstances_aggravantes" in details:
            for example in details["circonstances_aggravantes"]:
                example_paragraph = doc.add_paragraph(style='List Bullet 2')
                example_paragraph.add_run(f"      {example}")

    # Ajouter l'Unité Administrative
    doc.add_heading('Unité Administrative', level=2)
    admin_unit = parameters["unité_administrative"]
    for key, value in admin_unit.items():
        doc.add_paragraph(f"{key.replace('_', ' ').capitalize()}: {value}", style='List Bullet')

    # Ajouter les Cycles Scolaires
    doc.add_heading('Cycles Scolaires', level=2)
    school_cycles = parameters.get("cycles_scolaires", {})
    age_ranges = school_cycles.get("tranches_d_age", [])
    doc.add_paragraph(f"Tranches d'âge: {age_ranges}", style='List Bullet')

    # Sauvegarder le document Word dans un objet BytesIO
    doc_output = BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)

    return doc_output
