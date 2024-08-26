import pandas as pd
#import fuzzywuzzy
from fuzzywuzzy import process
import numpy as np
import datetime
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Border, Side, Font, Alignment
from openpyxl.cell.cell import MergedCell  # Import MergedCell
import docx
from docx.shared import Pt, RGBColor
import matplotlib.pyplot as plt
from docx.shared import Inches
from io import BytesIO


int_2 = '2.0'
int_3 = '3.0'
int_4 = '4.0'
int_5 = '5.0'
label_perc2 = '% 1-2'
label_perc3 = '% 3'
label_perc4 = '% 4'
label_perc5 = '% 5'
label_tot2 = '# 1-2'
label_tot3 = '# 3'
label_tot4 = '# 4'
label_tot5 = '# 5'
label_perc_tot = '% Tot PiN (3+)'
label_tot = '# Tot PiN (3+)'
label_admin_severity = 'Area severity'
label_tot_population = 'TotN'

int_acc = 'access'
int_agg= 'aggravating circumstances'
int_lc = 'learning condition'
int_penv = 'protected environment'
int_out = 'not falling within the PiN dimensions'
label_perc_acc = '% Access'
label_perc_agg= '% Aggravating circumstances'
label_perc_lc = '% Learning conditions'
label_perc_penv = '% Protected environment'
label_perc_out = '% Not falling within the PiN dimensions'
label_tot_acc = '# Access'
label_tot_agg= '# Aggravating circumstances'
label_tot_lc = '# Learning conditions'
label_tot_penv = '# Protected environment'
label_tot_out = '# Not falling within the PiN dimensions'

label_dimension_perc_tot = '% Tot in PiN Dimensions'
label_dimension_tot = '# Tot in PiN Dimensions'

label_dimension_tot_population = 'TotN'



# Define the colors
colors = {
    "light_beige": "#FFF2CC",
    "light_orange": "#F4B183",
    "dark_orange": "#ED7D31",
    "darker_orange": "#C65911",
    "light_blue": "#DDEBF7",
    "light_pink": "#b3b389",
    "light_yellow": "#ffffc5",
    "white": "#FFFFFF",
    "bluepin": "#004bb4",
    'gray': '#e0e0e0',
    'stratagray': '#F0F0F0'
}
# Define the columns to color
color_mapping = {
    label_perc2: colors["light_beige"],
    label_tot2: colors["light_beige"],
    label_perc3: colors["light_orange"],
    label_tot3: colors["light_orange"],
    label_perc4: colors["dark_orange"],
    label_tot4: colors["dark_orange"],
    label_perc5: colors["darker_orange"],
    label_tot5: colors["darker_orange"],
    label_perc_tot: colors["light_blue"],
    label_admin_severity: colors["light_blue"],
    label_tot: colors["light_blue"]
}
def format_number(num):
    """Convert a large number to a shorter format with K or M."""
    if num >= 1_000_000:
        return f"{num / 1_000_000:.1f}M"
    elif num >= 1_000:
        return f"{num / 1_000:.0f}K"
    else:
        return f"{num:.0f}"
    


def create_snapshot_PiN(country_label, final_overview_df, final_overview_dimension_df):


    country_name = country_label.split('__')[0]  # Extract the part before the "__"

    # Initialize the Word document
    doc = docx.Document()

    # Main title style
    title = doc.add_paragraph(f'{country_name} – PiN snapshot')
    title_run = title.runs[0]
    title_run.font.size = Pt(24)  # Make the title larger
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title.alignment = 1  # Center the title

    # Add some spacing after the title
    doc.add_paragraph()

    # Session 1: Children in Need (5-17 y.o.)
    section1 = doc.add_heading('Children in Need (5-17 y.o.)', level=2)
    section1_run = section1.runs[0]
    section1_run.font.size = Pt(20)  # Customize the section header size
    section1_run.bold = True
    section1_run.font.name = 'Calibri'
    section1.alignment = 0  # Left align
    doc.add_paragraph()


    # Create a table with one row and two columns to hold the text and the pie chart side by side
    table = doc.add_table(rows=1, cols=2)
    cell_width1 = Inches(3)  # Adjust this value to control the width of the cells
    cell_width2 = Inches(4)  # Adjust this value to control the width of the cells

    #table.autofit = True

    # Retrieve data for the total row
    row_tot = final_overview_df.loc[final_overview_df['Strata'] == 'TOTAL (5-17 y.o.)']
    total_population = row_tot[label_tot_population].values[0]
    percentage_in_need = row_tot[label_perc_tot].values[0]
    tot_in_need = row_tot[label_tot].values[0]
    percentage_2 = row_tot[label_perc2].values[0]
    percentage_3 = row_tot[label_perc3].values[0]
    percentage_4 = row_tot[label_perc4].values[0]
    percentage_5 = row_tot[label_perc5].values[0]

    # Add text before the pie chart
    cell_text = table.cell(0, 0)
    paragraph_text = cell_text.paragraphs[0]
    run_text = paragraph_text.add_run(f"{percentage_in_need}% amounts to {format_number(tot_in_need)} children in Need")
    run_text.font.size = Pt(15)
    run_text.font.color.rgb = RGBColor(5, 77, 180)  # Blue text
    run_text.bold = True

    cell_chart = table.cell(0, 1)
    cell_chart.width = cell_width1


    ####################### PIE ###############################
    # Generate the pie chart
    labels_pie = ['severity 1-2', 'severity 3', 'severity 4', 'severity 5']
    sizes = [percentage_2, percentage_3, percentage_4, percentage_5]
    colors_pie = [color_mapping[label_perc2], color_mapping[label_perc3], color_mapping[label_perc4], color_mapping[label_perc5]]

    # Create the pie chart
    fig, ax = plt.subplots(figsize=(3, 3))  # Smaller figure size
    wedges, texts, autotexts = ax.pie(sizes, labels=None, autopct='%1.1f%%', startangle=90, colors=colors_pie)

    # Style the chart
    plt.setp(autotexts, size=10, color="black")

    # Adjust label positions to avoid overlap by manually tweaking positions for smaller segments
    for i, autotext in enumerate(autotexts):
        if sizes[i] < 4:  # Apply special positioning for smaller slices
            position = autotext.get_position()
            autotext.set_position((position[0] * 1.5, position[1] * 1.5))  # Move the text further out
        else:
            position = autotext.get_position()
            autotext.set_position((position[0] * 1.05, position[1] * 1.05))  # Slightly adjust larger slices

    ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

    # Add a title to the pie chart
    title_pie = f"Children 5-17 y.o. -- ToT {format_number(total_population)}"
    plt.title(title_pie, fontsize=12)

    # Add a legend below the chart
    plt.legend(wedges, labels_pie, loc="lower center", bbox_to_anchor=(0.5, -0.1), ncol=4, fontsize=8)  # Smaller font size for legend
    # Save the pie chart
    pie_chart_jpeg_path = "pie_chart_with_text_v3.jpeg"
    plt.savefig(pie_chart_jpeg_path, format='jpeg', bbox_inches='tight', dpi=300)  # Higher DPI for better quality
    plt.close(fig)

    # Insert the pie chart into the second column of the table
    run_chart = cell_chart.paragraphs[0].add_run()
    run_chart.add_picture(pie_chart_jpeg_path, width=cell_width2)  # Adjust the width to match the cell width



    # Section: Population groups
    doc.add_heading('Population groups', level=2)

    excluded_strata = ['TOTAL (5-17 y.o.)',
        "Girls", "Boys", "Female (MSNA)", "Male (MSNA)", "ECE (5 y.o.)", 
        "Primary school", "Upper primary school", "Secondary school", "Children with disability"
    ]

    groups = []
    severity_1_2_groups = []
    severity_3_groups = []
    severity_4_groups = []
    severity_5_groups = []
    # Iterate over the DataFrame rows to create the content
    for _, row_pop in final_overview_df.iterrows():
        population_group = row_pop['Population group']
        total_population_in_need = row_pop[label_tot]
        perc_population_in_need = row_pop[label_perc_tot]
        strata = row_pop['Strata']
        percentage_2_pop = row_pop[label_perc2]
        percentage_3_pop = row_pop[label_perc3]
        percentage_4_pop = row_pop[label_perc4]
        percentage_5_pop = row_pop[label_perc5]
        
        if strata not in excluded_strata:
            # For excluded strata, only include the percentages and the total percentage in need
            doc.add_paragraph(
                f"{population_group}: {perc_population_in_need}% are in need, which amounts to {format_number(total_population_in_need)} children.",
                style='List Bullet'
            )
            # Append values to the lists for plotting
            groups.append(population_group)
            severity_1_2_groups.append(percentage_2_pop)
            severity_3_groups.append(percentage_3_pop)
            severity_4_groups.append(percentage_4_pop)
            severity_5_groups.append(percentage_5_pop)

    ind = np.arange(len(groups))
    width = 0.5

    # Create the bar chart
    fig, ax = plt.subplots(figsize=(5, 3))  # Smaller figure size
    p1 = ax.barh(ind, severity_1_2_groups, width, color=color_mapping[label_perc2], label='severity 1-2')
    p2 = ax.barh(ind, severity_3_groups, width, left=severity_1_2_groups, color=color_mapping[label_perc3], label='severity 3')
    p3 = ax.barh(ind, severity_4_groups, width, left=np.array(severity_1_2_groups) + np.array(severity_3_groups), color=color_mapping[label_perc4], label='severity 4')
    p4 = ax.barh(ind, severity_5_groups, width, left=np.array(severity_1_2_groups) + np.array(severity_3_groups) + np.array(severity_4_groups), color=color_mapping[label_perc5], label='severity 5')

    ax.set_xlabel('Percentage')
    ax.set_title('Severity distribution -- Population group')
    ax.xaxis.grid(True, linestyle='--', which='major', color='gray', alpha=0.7)

    ax.set_yticks(ind)
    ax.set_yticklabels(groups)
    ax.legend(loc="center left", bbox_to_anchor=(1, 0.5))  # Move legend to the right

    # Save the bar chart
    bar_chart_path = "bar_chart_with_text.jpeg"
    plt.savefig(bar_chart_path, format='jpeg', bbox_inches='tight', dpi=300)  # Higher DPI for better quality
    plt.close(fig)
    # Insert the bar chart after the corresponding section
    doc.add_picture(bar_chart_path)

    # Section: Profile of children in need
    doc.add_heading('Profile of children in need', level=2)
    doc.add_paragraph(
        "This section provides an overview of different age groups and specific vulnerable groups of children in need."
    )

    # Retrieve data for the total row
    row_girl = final_overview_df.loc[final_overview_df['Strata'] == 'Female (MSNA)']
    total_population_girl = row_girl[label_tot_population].values[0]
    percentage_in_need_girl = row_girl[label_perc_tot].values[0]
    tot_in_need_girl = row_girl[label_tot].values[0]
    percentage_2_girl = row_girl[label_perc2].values[0]
    percentage_3_girl = row_girl[label_perc3].values[0]
    percentage_4_girl = row_girl[label_perc4].values[0]
    percentage_5_girl = row_girl[label_perc5].values[0]

    ####################### PIE ###############################
    # Generate the pie chart
    sizes_girl = [percentage_2_girl, percentage_3_girl, percentage_4_girl, percentage_5_girl]
    # Create the pie chart
    fig_girl, ax = plt.subplots(figsize=(3.5, 3.5)) 
    wedges, texts, autotexts = ax.pie(sizes_girl, labels=None, autopct='%1.2f%%', startangle=90, colors=colors_pie)

    # Style the chart
    plt.setp(autotexts, size=10, color="black")

    # Adjust label positions to avoid overlap by manually tweaking positions for smaller segments
    for i, autotext in enumerate(autotexts):
        if sizes_girl[i] < 4:  # Apply special positioning for smaller slices
            position = autotext.get_position()
            autotext.set_position((position[0] * 1.5, position[1] * 1.5))  # Move the text further out
        else:
            position = autotext.get_position()
            autotext.set_position((position[0] * 1.05, position[1] * 1.05))  # Slightly adjust larger slices

    ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

    # Add a title to the pie chart
    title_pie = "Girls, severity distribution"
    plt.title(title_pie, fontsize=12)

    # Add a legend below the chart
    plt.legend(wedges, labels_pie, loc="lower center", bbox_to_anchor=(0.5, -0.1), ncol=4)
    # Save the pie chart
    pie_chart_jpeg_path_girl = "pie_chart_with_text_girl.jpeg"
    plt.savefig(pie_chart_jpeg_path_girl, format='jpeg', bbox_inches='tight', dpi=300)  # Higher DPI for better quality
    plt.close(fig_girl)

    # Insert the pie chart into the Word document
    #doc.add_picture(pie_chart_jpeg_path_girl)


    # Retrieve data for the total row
    row_boy = final_overview_df.loc[final_overview_df['Strata'] == 'Male (MSNA)']
    total_population_boy = row_boy[label_tot_population].values[0]
    percentage_in_need_boy = row_boy[label_perc_tot].values[0]
    tot_in_need_boy = row_boy[label_tot].values[0]
    percentage_2_boy = row_boy[label_perc2].values[0]
    percentage_3_boy = row_boy[label_perc3].values[0]
    percentage_4_boy = row_boy[label_perc4].values[0]
    percentage_5_boy = row_boy[label_perc5].values[0]

    ####################### PIE ###############################
    # Generate the pie chart
    sizes_boy = [percentage_2_boy, percentage_3_boy, percentage_4_boy, percentage_5_boy]
    fig_boy, ax = plt.subplots(figsize=(3.5, 3.5))  # Smaller figure size
    wedges, texts, autotexts = ax.pie(sizes_boy, labels=None, autopct='%1.2f%%', startangle=90, colors=colors_pie)

    # Style the chart
    plt.setp(autotexts, size=10,  color="black")

    # Adjust label positions to avoid overlap by manually tweaking positions for smaller segments
    for i, autotext in enumerate(autotexts):
        if sizes_boy[i] < 4:  # Apply special positioning for smaller slices
            position = autotext.get_position()
            autotext.set_position((position[0] * 1.5, position[1] * 1.5))  # Move the text further out
        else:
            position = autotext.get_position()
            autotext.set_position((position[0] * 1.05, position[1] * 1.05))  # Slightly adjust larger slices

    ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

    # Add a title to the pie chart
    title_pie = "Boys, severity distribution"
    plt.title(title_pie, fontsize=12)

    # Add a legend below the chart
    plt.legend(wedges, labels_pie, loc="lower center", bbox_to_anchor=(0.5, -0.1), ncol=4)
    # Save the pie chart
    pie_chart_jpeg_path_boy = "pie_chart_with_text_boy.jpeg"
    plt.savefig(pie_chart_jpeg_path_boy, format='jpeg', bbox_inches='tight', dpi=300)  # Higher DPI for better quality
    plt.close(fig_boy)

    # Insert the pie chart into the Word document
    #doc.add_picture(pie_chart_jpeg_path_boy)

    # Create a table with one row and two columns
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # Insert the first image (Girls pie chart) into the first cell
    cell_girl = table.cell(0, 0)
    paragraph_girl = cell_girl.paragraphs[0]
    run_girl = paragraph_girl.add_run()
    run_girl.add_picture(pie_chart_jpeg_path_girl, width=Inches(3.5))  # Adjust the width as necessary

    # Insert the second image (Boys pie chart) into the second cell
    cell_boy = table.cell(0, 1)
    paragraph_boy = cell_boy.paragraphs[0]
    run_boy = paragraph_boy.add_run()
    run_boy.add_picture(pie_chart_jpeg_path_boy, width=Inches(3.5))  # Adjust the width as necessary


    # Retrieve data for various population groups
    row_ece = final_overview_df.loc[final_overview_df['Strata'] == 'ECE (5 y.o.)']
    row_primary = final_overview_df.loc[final_overview_df['Strata'] == 'Primary school']
    row_upper_primary = final_overview_df.loc[final_overview_df['Strata'] == 'Upper primary school']
    row_upper_primary = row_upper_primary if not row_upper_primary.empty else None
    row_secondary = final_overview_df.loc[final_overview_df['Strata'] == 'Secondary school']
    row_disability = final_overview_df.loc[final_overview_df['Strata'] == 'Children with disability']

    tot_in_need_ece = row_ece[label_tot].values[0]
    tot_in_need_primary = row_primary[label_tot].values[0]
    tot_in_need_upper_primary = row_upper_primary[label_tot].values[0] if row_upper_primary is not None else None
    tot_in_need_secondary = row_secondary[label_tot].values[0]
    tot_in_need_disability = row_disability[label_tot].values[0]
    doc.add_paragraph()
    doc.add_paragraph(
        "The numbers below reflect the estimated total number of children in need in each category."
    )
    # Create the profile_groups list conditionally
    profile_groups = [
        f"5 y.o. children in need: {format_number(tot_in_need_ece)}",
        f"Primary school-aged children in need: {format_number(tot_in_need_primary)}",
        f"Secondary school-aged children in need: {format_number(tot_in_need_secondary)}",
        f"Children with disability in need: {format_number(tot_in_need_disability)}"
    ]
    # Add Upper Primary school data only if it exists
    if tot_in_need_upper_primary is not None:
        profile_groups.insert(2, f"Upper Primary school-aged children in need: {format_number(tot_in_need_upper_primary)}")

    profile_text = "\n".join(profile_groups)
    doc.add_paragraph(profile_text)


    doc.add_page_break()

    # Section: Profile of Needs in Myanmar
    section2 = doc.add_heading(f'Profile of Needs in {country_name}', level=2)
    section2_run = section2.runs[0]
    section2_run.font.size = Pt(20)  # Customize the section header size
    section2_run.bold = True
    section2_run.font.name = 'Calibri'
    section2.alignment = 0  # Left align
    # Add a line of spacing after the section heading
    doc.add_paragraph()

    # Retrieve data for the total row
    row_tot_dimension = final_overview_dimension_df.loc[final_overview_dimension_df['Strata'] == 'TOTAL (5-17 y.o.)']
    percentage_access = row_tot_dimension[label_perc_acc].values[0]
    percentage_ag = row_tot_dimension[label_perc_agg].values[0]
    percentage_lc = row_tot_dimension[label_perc_lc].values[0]
    percentage_pe = row_tot_dimension[label_perc_penv].values[0]


    doc.add_paragraph(
        "By examining the four main dimensions of need—Access, Learning Conditions, Protected Environment, "
        "and Aggravating Circumstances—we can identify the most common needs and better understand the support "
        "required for children in need. Among children aged 5 to 17 years:\n"
        f"{percentage_access}% lack access to school\n"
        f"{percentage_ag}% are affected by aggravating circumstances\n"
        f"{percentage_lc}% experience poor learning conditions\n"
        f"{percentage_pe}% attend school in unsafe environments.\n"
    )



    # Section: Variation in needs across different population groups
    doc.add_paragraph("The variation in needs across different population groups is noteworthy:")


    # Iterate over the DataFrame rows to create the content
    for _, row_pop_dim in final_overview_dimension_df.iterrows():
        population_group = row_pop_dim['Population group']
        strata = row_pop_dim['Strata']
        percentage_acc_pop = row_pop_dim[label_perc_acc]
        percentage_ag_pop = row_pop_dim[label_perc_agg]
        percentage_lc_pop = row_pop_dim[label_perc_lc]
        percentage_pe_pop = row_pop_dim[label_perc_penv]
        
        if strata not in excluded_strata:
            # For excluded strata, only include the percentages and the total percentage in need
            doc.add_paragraph(
                f"{population_group}: {percentage_acc_pop}% lack access to school, {percentage_acc_pop}% face aggravating circumstances, {percentage_lc_pop}% experience poor learning conditions, and {percentage_pe_pop}% attend school in unsafe environments.",
                style='List Bullet'
            )

    doc.add_paragraph("This data highlights the differing educational needs among these vulnerable groups.")

    # Save the Word document to a BytesIO object
    doc_output = BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)

    # Return the BytesIO object containing the Word document
    return doc_output
